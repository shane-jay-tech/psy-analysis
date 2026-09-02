"""数据加载入口测试（v3.7 N7：jsPsych JSON/JSONL）。"""

import io
import json

import pandas as pd
import pytest

from src.data.loader import load_data, load_jspsych_json


def _make_jspsych_json_array() -> bytes:
    trials = [
        {"trial_index": 0, "trial_type": "html-keyboard-response", "rt": 412, "response": "f"},
        {"trial_index": 1, "trial_type": "html-keyboard-response", "rt": 388, "response": "j"},
        {"trial_index": 2, "trial_type": "survey-text", "rt": 9000, "response": '{"Q1":"焦虑"}'},
    ]
    return json.dumps(trials, ensure_ascii=False).encode("utf-8")


def _make_jspsych_jsonl() -> bytes:
    lines = [
        '{"trial_index": 0, "trial_type": "fixation", "rt": null}',
        '{"trial_index": 1, "trial_type": "stimulus", "rt": 520, "correct": true}',
        '{"trial_index": 2, "trial_type": "stimulus", "rt": 480, "correct": false}',
    ]
    return ("\n".join(lines)).encode("utf-8")


class TestJsPsychJsonLoader:
    def test_load_json_array(self):
        buf = io.BytesIO(_make_jspsych_json_array())
        df, meta = load_jspsych_json(buf)
        assert len(df) == 3
        assert meta["source_type"] == "jspsych_json"
        assert meta["format"] == "json_array"
        assert meta["row_count"] == 3
        # column normalization 应至少触发一个映射
        assert "trial_type" in df.columns or "试次类型" in df.columns

    def test_load_jsonl(self):
        buf = io.BytesIO(_make_jspsych_jsonl())
        df, meta = load_jspsych_json(buf)
        assert len(df) == 3
        assert meta["format"] == "jsonl"

    def test_empty_raises(self):
        with pytest.raises(ValueError, match="为空"):
            load_jspsych_json(io.BytesIO(b""))

    def test_invalid_array_raises(self):
        with pytest.raises(ValueError, match="解析失败"):
            load_jspsych_json(io.BytesIO(b"[bad json}"))

    def test_single_object_treated_as_one_jsonl_line(self):
        # `{"foo":1}` 不是数组但是一个合法 JSONL 行 → 容错为 1 行
        df, meta = load_jspsych_json(io.BytesIO(b'{"foo": 1}'))
        assert len(df) == 1
        assert meta["format"] == "jsonl"

    def test_jsonl_skips_bad_lines(self):
        bad = b'{"a":1}\nNOT JSON\n{"a":2}\n'
        df, meta = load_jspsych_json(io.BytesIO(bad))
        # 跳过坏行，剩 2 条
        assert len(df) == 2
        assert meta["format"] == "jsonl"

    def test_all_bad_lines_raises(self):
        bad = b"only bad lines\nstill bad\n"
        with pytest.raises(ValueError, match="未能从"):
            load_jspsych_json(io.BytesIO(bad))


class TestLoadDataDispatchesJson:
    """load_data 应根据 .json/.jsonl 扩展名路由到 jsPsych 解析。"""

    def test_load_data_routes_json(self, tmp_path):
        path = tmp_path / "exp.json"
        path.write_bytes(_make_jspsych_json_array())
        df, meta = load_data(str(path))
        assert meta["source_type"] == "jspsych_json"
        assert len(df) == 3

    def test_load_data_routes_jsonl(self, tmp_path):
        path = tmp_path / "exp.jsonl"
        path.write_bytes(_make_jspsych_jsonl())
        df, meta = load_data(str(path))
        assert meta["format"] == "jsonl"

    def test_unsupported_extension_message_includes_jspsych(self, tmp_path):
        path = tmp_path / "data.txt"
        path.write_text("hello", encoding="utf-8")
        with pytest.raises(ValueError, match="jsPsych"):
            load_data(str(path))

    def test_load_data_usecols_filter(self, tmp_path):
        path = tmp_path / "exp.json"
        path.write_bytes(_make_jspsych_json_array())
        # 列名归一化后 trial_index → 试次序号；usecols 用归一化后的名字
        df, meta = load_data(str(path), usecols=["试次序号"])
        assert list(df.columns) == ["试次序号"]
        assert meta["col_count"] == 1


# ---------------------------------------------------------------------------
# v3.7 N7+: Word (.docx) 表格加载
# ---------------------------------------------------------------------------

def _make_word_with_table(path, rows):
    from docx import Document
    doc = Document()
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = str(val)
    doc.save(str(path))


class TestWordLoader:
    from src.data.loader import load_word_table

    def test_load_word_table_basic(self, tmp_path):
        from src.data.loader import load_word_table
        path = tmp_path / "data.docx"
        _make_word_with_table(path, [
            ["id", "score", "group"],
            ["s01", "85", "A"],
            ["s02", "92", "B"],
            ["s03", "78", "A"],
        ])
        df, meta = load_word_table(str(path))
        assert meta["source_type"] == "word"
        assert meta["row_count"] == 3
        assert list(df.columns) == ["id", "score", "group"]
        # 数值列应被自动转 numeric
        assert pd.api.types.is_numeric_dtype(df["score"])

    def test_word_no_table_raises(self, tmp_path):
        from docx import Document
        from src.data.loader import load_word_table
        path = tmp_path / "empty.docx"
        Document().save(str(path))
        with pytest.raises(ValueError, match="未找到表格"):
            load_word_table(str(path))

    def test_load_data_routes_docx(self, tmp_path):
        path = tmp_path / "study.docx"
        _make_word_with_table(path, [
            ["item", "value"],
            ["a", "1"],
            ["b", "2"],
        ])
        df, meta = load_data(str(path))
        assert meta["source_type"] == "word"
        assert len(df) == 2

    def test_word_table_index_out_of_range(self, tmp_path):
        from src.data.loader import load_word_table
        path = tmp_path / "data.docx"
        _make_word_with_table(path, [["a"], ["1"]])
        with pytest.raises(ValueError, match="索引越界"):
            load_word_table(str(path), table_index=5)


# ---------------------------------------------------------------------------
# v3.7 N7+: Markdown (.md) 表格加载
# ---------------------------------------------------------------------------

class TestMarkdownLoader:
    def test_basic_gfm_table(self, tmp_path):
        from src.data.loader import load_markdown_table
        path = tmp_path / "study.md"
        path.write_text(
            "# 数据\n\n"
            "| id  | rt  | correct |\n"
            "|-----|-----|---------|\n"
            "| s01 | 412 | 1       |\n"
            "| s02 | 388 | 0       |\n"
            "| s03 | 521 | 1       |\n",
            encoding="utf-8",
        )
        df, meta = load_markdown_table(str(path))
        assert meta["source_type"] == "markdown"
        assert meta["n_tables"] == 1
        assert list(df.columns) == ["id", "rt", "correct"]
        assert len(df) == 3
        assert pd.api.types.is_numeric_dtype(df["rt"])

    def test_no_table_raises(self, tmp_path):
        from src.data.loader import load_markdown_table
        path = tmp_path / "doc.md"
        path.write_text("# 标题\n\n这只是一段普通文本。", encoding="utf-8")
        with pytest.raises(ValueError, match="未找到表格"):
            load_markdown_table(str(path))

    def test_two_tables_select_by_index(self, tmp_path):
        from src.data.loader import load_markdown_table
        path = tmp_path / "doc.md"
        path.write_text(
            "| a | b |\n|---|---|\n| 1 | 2 |\n| 3 | 4 |\n\n"
            "正文段落\n\n"
            "| x | y |\n|---|---|\n| 9 | 8 |\n",
            encoding="utf-8",
        )
        from src.data.loader import load_markdown_table
        df0, _ = load_markdown_table(str(path), table_index=0)
        df1, meta1 = load_markdown_table(str(path), table_index=1)
        assert list(df0.columns) == ["a", "b"]
        assert list(df1.columns) == ["x", "y"]
        assert meta1["n_tables"] == 2

    def test_load_data_routes_md(self, tmp_path):
        path = tmp_path / "study.md"
        path.write_text(
            "| id | val |\n|---|---|\n| a | 1 |\n| b | 2 |\n",
            encoding="utf-8",
        )
        df, meta = load_data(str(path))
        assert meta["source_type"] == "markdown"
        assert len(df) == 2

    def test_load_data_routes_markdown_extension(self, tmp_path):
        path = tmp_path / "study.markdown"
        path.write_text(
            "| q | a |\n|---|---|\n| 1 | 2 |\n",
            encoding="utf-8",
        )
        df, meta = load_data(str(path))
        assert meta["source_type"] == "markdown"

    def test_unsupported_extension_lists_word_and_markdown(self, tmp_path):
        path = tmp_path / "data.txt"
        path.write_text("nope", encoding="utf-8")
        with pytest.raises(ValueError, match="Word.*Markdown|jsPsych"):
            load_data(str(path))


class TestPivotJspsychToWide:
    """v3.9 N9: jsPsych 长表 → 被试级宽表。"""

    def _make_long_df(self):
        return pd.DataFrame({
            "subject": ["s1", "s1", "s2", "s2", "s3", "s3"],
            "condition": ["congruent", "incongruent"] * 3,
            "rt": [400, 600, 380, 580, 410, 620],
        })

    def test_basic_pivot_with_explicit_cols(self):
        from src.data.loader import pivot_jspsych_to_wide
        wide, meta = pivot_jspsych_to_wide(
            self._make_long_df(),
            subject_col="subject",
            condition_col="condition",
            value_col="rt",
        )
        assert wide.shape == (3, 3)  # 3 subjects, 1 subject col + 2 condition cols
        assert "congruent" in wide.columns
        assert "incongruent" in wide.columns
        assert meta["n_subjects"] == 3
        assert meta["n_conditions"] == 2
        assert meta["agg"] == "mean"

    def test_pivot_aggregates_multiple_trials_per_cell_as_mean(self):
        from src.data.loader import pivot_jspsych_to_wide
        df = pd.DataFrame({
            "subject": ["s1", "s1", "s1"],
            "condition": ["a", "a", "a"],
            "rt": [100, 200, 300],
        })
        wide, _ = pivot_jspsych_to_wide(df, value_col="rt")
        assert wide["a"].iloc[0] == 200  # mean

    def test_pivot_median_aggregation(self):
        from src.data.loader import pivot_jspsych_to_wide
        df = pd.DataFrame({
            "subject": ["s1", "s1", "s1", "s1"],
            "condition": ["a", "a", "a", "a"],
            "rt": [100, 200, 300, 400],
        })
        wide, meta = pivot_jspsych_to_wide(df, value_col="rt", agg="median")
        assert wide["a"].iloc[0] == 250
        assert meta["agg"] == "median"

    def test_pivot_auto_resolves_chinese_value_col(self):
        from src.data.loader import pivot_jspsych_to_wide
        df = pd.DataFrame({
            "subject": ["s1", "s1"],
            "condition": ["a", "b"],
            "反应时_ms": [400, 600],
        })
        wide, meta = pivot_jspsych_to_wide(df)  # 默认 value_col="反应时_ms"
        assert "a" in wide.columns
        assert "b" in wide.columns
        assert meta["value_col"] == "反应时_ms"

    def test_pivot_auto_resolves_subject_id_variant(self):
        from src.data.loader import pivot_jspsych_to_wide
        df = pd.DataFrame({
            "subj_id": ["s1", "s2"],
            "condition": ["a", "a"],
            "rt": [400, 500],
        })
        wide, meta = pivot_jspsych_to_wide(df, value_col="rt")
        assert meta["subject_col"] == "subj_id"
        assert wide.shape[0] == 2

    def test_pivot_raises_on_missing_subject(self):
        from src.data.loader import pivot_jspsych_to_wide
        df = pd.DataFrame({
            "trial_type": ["a", "b"],
            "rt": [100, 200],
        })
        with pytest.raises(ValueError, match="subject"):
            pivot_jspsych_to_wide(df, value_col="rt")

    def test_pivot_raises_on_empty_df(self):
        from src.data.loader import pivot_jspsych_to_wide
        with pytest.raises(ValueError, match="为空"):
            pivot_jspsych_to_wide(pd.DataFrame())

    def test_pivot_meta_pivoted_from_marker(self):
        from src.data.loader import pivot_jspsych_to_wide
        wide, meta = pivot_jspsych_to_wide(
            self._make_long_df(), value_col="rt"
        )
        assert meta["pivoted_from"] == "jspsych_long"
        assert meta["source_type"] == "jspsych_pivoted"

class _NamedBytesIO(io.BytesIO):
    def __init__(self, data: bytes, name: str):
        super().__init__(data)
        self.name = name


class TestExcelLoaderPandas3Regression:
    """v5.8: pandas 3.x 下 sheet_name=None 返回 dict 导致 load_excel 崩溃的回归测试。"""

    @staticmethod
    def _make_xlsx_bytes(sheets: list[tuple[str, list[list]]]) -> bytes:
        import openpyxl
        wb = openpyxl.Workbook()
        wb.remove(wb.active)
        for title, rows in sheets:
            ws = wb.create_sheet(title=title)
            for row in rows:
                ws.append(row)
        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    def test_single_sheet_returns_dataframe(self):
        data = self._make_xlsx_bytes([("Sheet1", [["x", "y"], [1, 10], [2, 20]])])
        df, meta = load_data(_NamedBytesIO(data, "single.xlsx"))
        assert isinstance(df, pd.DataFrame)
        assert list(df.columns) == ["x", "y"]
        assert len(df) == 2
        assert meta["source_type"] == "excel"
        assert meta["sheet_name"] == "Sheet1"

    def test_multi_sheet_takes_first(self):
        data = self._make_xlsx_bytes([
            ("第一表", [["a"], [1]]),
            ("第二表", [["b"], [2]]),
        ])
        df, meta = load_data(_NamedBytesIO(data, "multi.xlsx"))
        assert meta["sheet_name"] == "第一表"
        assert list(df.columns) == ["a"]

    def test_explicit_sheet_name(self):
        data = self._make_xlsx_bytes([
            ("第一表", [["a"], [1]]),
            ("第二表", [["b"], [2]]),
        ])
        df, meta = load_data(_NamedBytesIO(data, "multi.xlsx"), sheet_name="第二表")
        assert list(df.columns) == ["b"]

    def test_consumed_pointer_still_loads(self):
        """v5.8: 大文件列预览（pd.read_csv(nrows=0)）消耗指针后 load_data 仍应完整加载。"""
        csv_bytes = ("a,b\n" + "\n".join(f"{i},{i * 2}" for i in range(500))).encode("utf-8")
        f = _NamedBytesIO(csv_bytes, "large.csv")
        pd.read_csv(f, nrows=0)  # 模拟 app.py 大文件预览
        df, meta = load_data(f)
        assert len(df) == 500
        assert meta["row_count"] == 500
