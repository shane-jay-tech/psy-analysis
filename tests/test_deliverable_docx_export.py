"""研究交付包 Word (.docx) 导出测试。"""

import pytest

from src.paper_writer.draft_bundle import PaperDraftBundle, PaperSection
from src.paper_writer.research_deliverable import ResearchDeliverableBundle
from src.output.docx_exporter import build_deliverable_docx


@pytest.fixture
def bundle_with_paper():
    paper = PaperDraftBundle(
        title="焦虑与自尊的关系研究",
        sections={
            "introduction": PaperSection(name="引言", markdown="焦虑是常见的情绪问题...", source="template"),
            "method": PaperSection(name="方法", markdown="## 被试\n大学生 N=200", source="template"),
            "result": PaperSection(name="结果", markdown="r=-.42, p<.001", source="data"),
        },
        source="template",
    )
    return ResearchDeliverableBundle(
        project_id="test_docx",
        title="焦虑与自尊的关系研究",
        paper_bundle=paper,
        analysis_cards=[{"method": "pearson_corr", "apa_text": "r(198) = -.42, p < .001"}],
        evidence_records=[{"citation_key": "wang2023", "claim": "焦虑与自尊负相关"}],
        data_cleaning_log=[{"step": "无效样本", "action": "剔除 3 例"}],
        method_recommendations=[{"recommendation": "Pearson 相关"}],
        health_report=[{"level": "WARN", "message": "样本量偏小"}],
    )


class TestDeliverableDocxExport:
    """验证 Word 导出生成有效 .docx 字节流。"""

    def test_basic_mode_generates_bytes(self, bundle_with_paper):
        result = build_deliverable_docx(bundle_with_paper, mode="basic")
        assert isinstance(result, bytes)
        assert len(result) > 1000
        assert result[:4] == b"PK\x03\x04"  # ZIP (docx) magic bytes

    def test_standard_mode_generates_bytes(self, bundle_with_paper):
        result = build_deliverable_docx(bundle_with_paper, mode="standard")
        assert isinstance(result, bytes)
        assert len(result) > len(build_deliverable_docx(bundle_with_paper, mode="basic"))

    def test_full_mode_generates_bytes(self, bundle_with_paper):
        result = build_deliverable_docx(bundle_with_paper, mode="full")
        assert isinstance(result, bytes)
        assert len(result) > len(build_deliverable_docx(bundle_with_paper, mode="standard"))

    def test_empty_bundle_still_works(self):
        bundle = ResearchDeliverableBundle(title="空项目")
        result = build_deliverable_docx(bundle, mode="basic")
        assert isinstance(result, bytes)
        assert result[:4] == b"PK\x03\x04"

    def test_docx_contains_title(self, bundle_with_paper):
        from docx import Document
        import io
        result = build_deliverable_docx(bundle_with_paper, mode="basic")
        doc = Document(io.BytesIO(result))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "焦虑与自尊" in all_text

    def test_standard_contains_evidence(self, bundle_with_paper):
        from docx import Document
        import io
        result = build_deliverable_docx(bundle_with_paper, mode="standard")
        doc = Document(io.BytesIO(result))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "wang2023" in all_text

    def test_full_contains_health_report(self, bundle_with_paper):
        from docx import Document
        import io
        result = build_deliverable_docx(bundle_with_paper, mode="full")
        doc = Document(io.BytesIO(result))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "样本量偏小" in all_text
