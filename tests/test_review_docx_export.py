"""Word 文献综述导出测试（build_review_docx）：返回可重开的 docx 且含标题与正文。"""

import io

from docx import Document

from src.output.docx_exporter import build_review_docx


def test_returns_bytes():
    data = build_review_docx("测试标题", "# 引言\n正文内容")
    assert isinstance(data, bytes)
    assert len(data) > 0


def test_reopenable():
    data = build_review_docx("标题", "一些正文")
    doc = Document(io.BytesIO(data))
    assert len(doc.paragraphs) > 0


def test_title_in_doc():
    data = build_review_docx("独特标题ABC123", "正文")
    doc = Document(io.BytesIO(data))
    assert any("独特标题ABC123" in p.text for p in doc.paragraphs)


def test_markdown_rendered():
    data = build_review_docx("标题", "## 主题一\n这里是正文内容段落")
    doc = Document(io.BytesIO(data))
    texts = [p.text for p in doc.paragraphs]
    assert any("主题一" in t for t in texts)
    assert any("正文内容段落" in t for t in texts)


def test_author_date_optional():
    data = build_review_docx("标题", "正文", author="张三", date="2026-06-02")
    doc = Document(io.BytesIO(data))
    texts = [p.text for p in doc.paragraphs]
    assert any("张三" in t for t in texts)


def test_empty_body_ok():
    data = build_review_docx("仅标题", "")
    doc = Document(io.BytesIO(data))
    assert any("仅标题" in p.text for p in doc.paragraphs)
