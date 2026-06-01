"""Word 论文导出器测试。"""

from __future__ import annotations

import io

import pandas as pd
import pytest
from docx import Document

from src.output.docx_exporter import (
    FigureItem, ThesisMeta, build_thesis_docx,
)


@pytest.fixture
def basic_meta():
    return ThesisMeta(
        title="社交焦虑与自尊的关系研究",
        author="张三",
        affiliation="某某大学心理学系",
        date="2026 年 5 月",
        abstract="本研究采用问卷调查法，以 200 名大学生为被试，"
                 "考察社交焦虑与自尊的关系。结果表明……",
        keywords=["社交焦虑", "自尊", "大学生"],
    )


@pytest.fixture
def desc_table():
    return pd.DataFrame({
        "变量": ["社交焦虑", "自尊"],
        "M": [42.31, 28.50],
        "SD": [8.12, 5.23],
        "n": [200, 200],
    })


@pytest.fixture
def small_png_bytes():
    """生成一个 100x100 的合法 PNG（python-docx 严格校验 PNG header）。"""
    try:
        from PIL import Image
    except ImportError:
        pytest.skip("Pillow 未安装，跳过图片嵌入测试")
    img = Image.new("RGB", (100, 100), color=(70, 130, 180))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _docx_to_text(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    return "\n".join(p.text for p in doc.paragraphs)


def test_minimum_meta_only(basic_meta):
    docx_bytes = build_thesis_docx(basic_meta)
    assert isinstance(docx_bytes, bytes)
    assert len(docx_bytes) > 1000  # docx zip 至少有这么大
    text = _docx_to_text(docx_bytes)
    assert basic_meta.title in text
    assert basic_meta.author in text


def test_method_and_result_markdown_rendered(basic_meta):
    method_md = "### 数据分析方法\n\n本研究采用 *t* 检验分析数据。"
    result_md = "### 结果\n\n描述统计结果显示，**M = 42.31**，*SD* = 8.12。"
    docx_bytes = build_thesis_docx(basic_meta, method_md=method_md, result_md=result_md)
    text = _docx_to_text(docx_bytes)
    assert "数据分析方法" in text
    assert "本研究采用" in text
    assert "结果" in text
    assert "描述统计" in text


def test_descriptive_table_inserted(basic_meta, desc_table):
    docx_bytes = build_thesis_docx(basic_meta, descriptive_table=desc_table)
    doc = Document(io.BytesIO(docx_bytes))
    assert len(doc.tables) == 1, "应该插入一张描述统计表"
    table = doc.tables[0]
    # 表头
    headers = [cell.text for cell in table.rows[0].cells]
    assert "变量" in headers
    assert "M" in headers
    # 数据
    body_text = "\n".join(cell.text for row in table.rows[1:] for cell in row.cells)
    assert "社交焦虑" in body_text
    assert "42.310" in body_text or "42.31" in body_text


def test_figure_embedded(basic_meta, small_png_bytes):
    fig = FigureItem(caption="社交焦虑分布图", png_bytes=small_png_bytes)
    docx_bytes = build_thesis_docx(basic_meta, figures=[fig])
    doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "社交焦虑分布图" in text
    # docx 图片实际嵌入需要遍历 inline shapes
    inline_shapes = doc.inline_shapes
    assert len(inline_shapes) == 1


def test_extra_tables_numbered_correctly(basic_meta, desc_table):
    extra = [
        ("各组均值比较", pd.DataFrame({"组别": ["A", "B"], "M": [10.5, 12.3]})),
    ]
    docx_bytes = build_thesis_docx(
        basic_meta, descriptive_table=desc_table, extra_tables=extra,
    )
    doc = Document(io.BytesIO(docx_bytes))
    paragraph_texts = [p.text for p in doc.paragraphs]
    assert any("表1" in t and "描述性统计" in t for t in paragraph_texts)
    assert any("表2" in t and "各组均值比较" in t for t in paragraph_texts)


def test_defense_qa_appendix_added(basic_meta):
    qa_md = "### Q1: 为什么用 t 检验？\n\n因为自变量是两组分类。"
    docx_bytes = build_thesis_docx(basic_meta, defense_qa_md=qa_md)
    text = _docx_to_text(docx_bytes)
    assert "答辩问题预演" in text
    assert "为什么用" in text


def test_plotly_figs_to_figure_items_handles_missing_kaleido(monkeypatch):
    """v2.8: kaleido 不可用时 plotly_figs_to_figure_items 应优雅返回空。"""
    from src.output.docx_exporter import plotly_figs_to_figure_items
    import src.visualization.paper_export as pe
    monkeypatch.setattr(pe, "_kaleido_available", lambda: False)

    import plotly.graph_objects as go
    fig = go.Figure(go.Scatter(x=[1, 2], y=[3, 4]))
    items = plotly_figs_to_figure_items([("test caption", fig)], palette="grayscale")
    assert items == []


def _kaleido_installed() -> bool:
    import importlib.util
    return importlib.util.find_spec("kaleido") is not None


@pytest.mark.skipif(not _kaleido_installed(), reason="kaleido not installed")
def test_plotly_figs_to_figure_items_smoke():
    """v2.8: 有 kaleido 时端到端转 PNG 成功。"""
    from src.output.docx_exporter import plotly_figs_to_figure_items
    import plotly.graph_objects as go
    fig = go.Figure(go.Bar(x=["A", "B"], y=[1, 2]))
    items = plotly_figs_to_figure_items(
        [("条件均值对比", fig)], palette="grayscale",
        width_px=600, height_px=400,
    )
    assert len(items) == 1
    assert items[0].caption == "条件均值对比"
    assert items[0].png_bytes[:8] == b"\x89PNG\r\n\x1a\n"


def test_figure_palette_param_doesnt_break_build(basic_meta, small_png_bytes):
    """build_thesis_docx 接受多种 figures 配色（已预生成的 PNG）。"""
    from src.output.docx_exporter import FigureItem
    figures = [FigureItem(caption="灰度图", png_bytes=small_png_bytes, width_cm=10.0)]
    docx_bytes = build_thesis_docx(basic_meta, figures=figures)
    assert docx_bytes.startswith(b"PK")  # docx 是 zip 格式


def test_custom_cover_template_appends_body(tmp_path, basic_meta):
    """v2.8: 用自定义封面模板时，模板内容保留，正文追加在后。"""
    from src.output.docx_exporter import build_thesis_with_custom_cover

    # 先创建一个简单的"封面模板"
    template_doc = Document()
    template_doc.add_heading("某某大学毕业论文封面", 0)
    template_doc.add_paragraph("学校代码: 12345")
    template_doc.add_paragraph("学院: 心理学院")
    template_path = tmp_path / "cover.docx"
    template_doc.save(template_path)

    docx_bytes = build_thesis_with_custom_cover(
        cover_template_path=str(template_path),
        meta=basic_meta,
        method_md="### 数据分析\n\n本研究采用 t 检验。",
        result_md="### 结果\n\n结果显示差异显著。",
    )
    assert isinstance(docx_bytes, bytes)
    assert docx_bytes.startswith(b"PK")

    # 读回检验
    out_doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in out_doc.paragraphs)
    # 模板内容应保留
    assert "毕业论文封面" in text
    assert "学校代码" in text
    # 正文也应被追加
    assert "数据分析" in text or "t 检验" in text


def test_custom_cover_template_invalid_path_raises_value_error(basic_meta, tmp_path):
    """v2.8: 模板路径无效时抛 ValueError。"""
    from src.output.docx_exporter import build_thesis_with_custom_cover
    bad_path = tmp_path / "does_not_exist.docx"
    with pytest.raises(ValueError) as exc_info:
        build_thesis_with_custom_cover(
            cover_template_path=str(bad_path),
            meta=basic_meta,
        )
    assert "封面模板" in str(exc_info.value) or ".docx" in str(exc_info.value)


def test_custom_cover_template_with_garbage_file_raises(basic_meta, tmp_path):
    """v2.8: 给个非 docx 文件应抛 ValueError，不崩溃。"""
    from src.output.docx_exporter import build_thesis_with_custom_cover
    bad_file = tmp_path / "fake.docx"
    bad_file.write_bytes(b"not a real docx file")
    with pytest.raises(ValueError):
        build_thesis_with_custom_cover(
            cover_template_path=str(bad_file),
            meta=basic_meta,
        )


def test_chinese_font_metadata_set(basic_meta):
    """确保中文字体被设置为东亚字体（防止 Word 显示乱码或英文字体）。"""
    docx_bytes = build_thesis_docx(basic_meta, method_md="### 方法\n\n中文段落测试。")
    doc = Document(io.BytesIO(docx_bytes))
    # 至少有一个 run 设置了 east-asia 字体
    found_cjk = False
    for para in doc.paragraphs:
        for run in para.runs:
            rpr = run._element.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr")
            if rpr is not None:
                rfonts = rpr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts")
                if rfonts is not None and rfonts.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"):
                    found_cjk = True
                    break
        if found_cjk:
            break
    assert found_cjk, "中文字体未被设置为 east-asia"
