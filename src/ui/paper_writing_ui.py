"""论文写作 UI — 从 app.py 拆分出的独立模块"""

import sys
import os

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

import streamlit as st
import pandas as pd

from src.paper_writer import PaperEngine
from src.paper_writer.literature_search_async import (
    search_literature_async,
    cancel_search_request,
    CancelledSearchError,
)
from src.paper_writer.defense_qa import (
    generate_defense_qa,
    group_qa_by_difficulty,
    render_qa_as_markdown,
)


from src.utils.i18n import t

section_order = [
    ("title", t("title")),
    ("abstract", t("abstract")),
    ("keywords", t("keywords")),
    ("introduction", t("introduction")),
    ("methods", t("methods")),
    ("results", t("results")),
    ("discussion", t("discussion")),
    ("references", t("references")),
]


def render_paper_writing_ui():
    """渲染论文写作界面"""
    # 初始化 PaperEngine
    if st.session_state.paper_engine is None:
        st.session_state.paper_engine = PaperEngine()

    engine = st.session_state.paper_engine

    st.title("📝 论文写作")
    st.caption("填写研究信息 → 导入分析结果 → 搜索文献 → 回答确认问题 → 生成初稿")

    # ---- 侧边栏 ----
    with st.sidebar:
        st.divider()
        st.header("📋 论文状态")

        # 信息缺口检查
        gaps = engine.get_gaps_summary()
        st.markdown(engine.get_gaps_summary())

        st.divider()
        st.header("💡 使用指南")
        st.markdown("""
        **写作流程：**
        1. 在"研究信息"标签页填写主题、假设、方法
        2. 在"分析结果"标签页导入已有分析结果（可选）
        3. 搜索相关文献或手动添加
        4. 回答系统提出的确认问题
        5. 生成完整论文初稿

        **格式标准：**
        - 参考文献格式：APA 7th Edition
        - 统计报告：APA7 规范
        - 论文结构：心理学报标准章节
        """)

        # 日志
        with st.expander("📜 操作日志"):
            for log in engine.get_logs():
                st.caption(log)

        # 重置
        if st.button("🔄 重置论文", width="stretch"):
            st.session_state.paper_engine = PaperEngine()
            st.rerun()

    # ---- 主区域：分步骤标签页 ----
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "1️⃣ 研究信息", "2️⃣ 文献管理", "3️⃣ 确认问题", "4️⃣ 生成论文", "5️⃣ 答辩模拟"
    ])

    # ============ 标签页1: 研究信息 ============
    with tab1:
        # 文件导入方法信息
        with st.expander("📥 从文件导入方法信息（JSON/Excel）", expanded=False):
            methods_file = st.file_uploader(
                "上传方法信息文件",
                type=["json", "xlsx", "xls"],
                key="methods_file_upload",
                help=(
                    "支持JSON和Excel格式。"
                    "JSON字段：participants_n, male_ratio, age_mean, age_sd, participants_desc, "
                    "materials, procedure, ethics, control_vars, theoretical, practical, limitations, future"
                ),
            )
            if methods_file is not None:
                import tempfile
                with tempfile.NamedTemporaryFile(delete=False, suffix=methods_file.name) as tmp:
                    tmp.write(methods_file.getvalue())
                    tmp_path = tmp.name
                try:
                    engine.import_methods_from_file(tmp_path)
                    st.success(f"已导入方法信息！")
                    st.rerun()
                except Exception as e:
                    st.error(f"导入失败：{e}")

        st.subheader("研究主题与假设")

        col_t1, col_t2 = st.columns(2)
        with col_t1:
            topic = st.text_input(
                "研究主题",
                value=engine.state.topic,
                placeholder="例如：大学生社交焦虑与自尊的关系研究",
                key="paper_topic",
            )
            title_hint = st.text_input(
                "论文标题（可选，留空则自动生成）",
                value=engine.state.title_hint,
                placeholder="不超过25字",
                key="paper_title_hint",
            )

            # 研究问题
            rq_text = st.text_area(
                "研究问题（每行一个）",
                value="\n".join(engine.state.research_questions),
                placeholder="例如：\n社交焦虑与自尊之间存在怎样的关系？\n社交焦虑是否存在性别差异？",
                height=100,
                key="paper_rq",
            )

        with col_t2:
            hyp_text = st.text_area(
                "研究假设（每行一个）",
                value="\n".join(engine.state.hypotheses),
                placeholder="例如：\nH1: 自尊负向预测社交焦虑\nH2: 社交焦虑存在显著的性别差异",
                height=150,
                key="paper_hyp",
            )

        if st.button("✅ 保存研究主题", key="save_topic"):
            rqs = [r.strip() for r in rq_text.split("\n") if r.strip()]
            hyps = [h.strip() for h in hyp_text.split("\n") if h.strip()]
            engine.set_topic(
                topic=topic.strip(),
                research_questions=rqs,
                hypotheses=hyps,
                title_hint=title_hint.strip(),
            )
            st.success("研究主题已保存！")
            st.rerun()

        st.divider()
        st.subheader("方法信息")

        col_m1, col_m2, col_m3 = st.columns(3)
        with col_m1:
            n = st.number_input("被试人数", min_value=0, value=engine.state.participants_n, key="paper_n")
            male_ratio = st.slider("男性比例", 0.0, 1.0, engine.state.male_ratio, 0.05, key="paper_male")
        with col_m2:
            age_mean = st.number_input("年龄 (M)", min_value=0.0, value=engine.state.age_mean, key="paper_age_m")
            age_sd = st.number_input("年龄 (SD)", min_value=0.0, value=engine.state.age_sd, key="paper_age_sd")
        with col_m3:
            participants_desc = st.text_area(
                "被试描述",
                value=engine.state.participants_desc,
                placeholder="例如：在校大学生，视力或矫正视力正常，无精神疾病史",
                key="paper_p_desc",
            )

        procedure = st.text_area(
            "施测程序",
            value=engine.state.procedure,
            placeholder="例如：采用网络问卷形式，通过Credamo平台发放，被试在阅读并签署知情同意书后开始作答，完成时间约15-20分钟。",
            key="paper_proc",
        )
        ethics = st.text_input(
            "伦理审批信息",
            value=engine.state.ethics,
            placeholder="例如：XX大学伦理委员会批准（批准号：XXX）",
            key="paper_ethics",
        )

        # 量表信息
        st.markdown("**测量工具（量表）**")
        n_materials = st.number_input("量表数量", min_value=0, max_value=10, value=max(1, len(engine.state.materials)), key="n_materials")
        materials = []
        for i in range(int(n_materials)):
            existing = engine.state.materials[i] if i < len(engine.state.materials) else {}
            col_mat1, col_mat2, col_mat3 = st.columns(3)
            with col_mat1:
                name = st.text_input(f"量表{i+1}名称", value=existing.get("name", ""), key=f"mat_name_{i}")
            with col_mat2:
                items = st.text_input(f"量表{i+1}题数", value=existing.get("items", ""), key=f"mat_items_{i}")
            with col_mat3:
                alpha = st.text_input(f"量表{i+1} α系数", value=existing.get("alpha", ""), key=f"mat_alpha_{i}")
            source = st.text_input(f"量表{i+1}来源", value=existing.get("source", ""), key=f"mat_src_{i}")
            if name:
                materials.append({"name": name, "items": items, "alpha": alpha, "source": source})

        # 讨论要点
        st.divider()
        st.subheader("讨论要点（可选）")
        col_d1, col_d2 = st.columns(2)
        with col_d1:
            theo = st.text_area(
                "理论贡献（每行一个）",
                value="\n".join(engine.state.theoretical_contributions),
                placeholder="例如：\n揭示了自尊对社交焦虑的影响机制\n拓展了XX理论在中文样本中的应用",
                key="paper_theo",
            )
            limit = st.text_area(
                "研究局限（每行一个）",
                value="\n".join(engine.state.limitations),
                placeholder="例如：\n横断设计无法推断因果关系\n样本代表性有限",
                key="paper_limit",
            )
        with col_d2:
            pract = st.text_area(
                "实践意义（每行一个）",
                value="\n".join(engine.state.practical_implications),
                placeholder="例如：\n为学校心理健康教育提供实证依据\n为大学生社交焦虑干预提供参考",
                key="paper_pract",
            )
            future = st.text_area(
                "未来方向（每行一个）",
                value="\n".join(engine.state.future_directions),
                placeholder="例如：\n采用纵向追踪设计验证因果关系\n扩大样本范围提高代表性",
                key="paper_future",
            )

        control_vars = st.text_input(
            "控制变量（用逗号分隔）",
            value=", ".join(engine.state.control_vars),
            placeholder="例如：性别, 年龄, 年级",
            key="paper_cvars",
        )

        if st.button("✅ 保存方法信息", key="save_methods", type="primary"):
            engine.set_methods(
                participants_n=int(n),
                male_ratio=male_ratio,
                age_mean=age_mean,
                age_sd=age_sd,
                participants_desc=participants_desc.strip(),
                materials=materials,
                procedure=procedure.strip(),
                ethics=ethics.strip(),
                control_vars=[c.strip() for c in control_vars.split(",") if c.strip()],
            )
            engine.set_discussion(
                theoretical=[t.strip() for t in theo.split("\n") if t.strip()],
                practical=[p.strip() for p in pract.split("\n") if p.strip()],
                limitations=[l.strip() for l in limit.split("\n") if l.strip()],
                future=[f.strip() for f in future.split("\n") if f.strip()],
            )
            st.success("方法信息已保存！")
            st.rerun()

    # ============ 标签页2: 文献管理 ============
    with tab2:
        # 文件导入文献
        with st.expander("📥 从文件导入文献（BibTeX / CSV / JSON）", expanded=False):
            lit_file = st.file_uploader(
                "上传文献文件",
                type=["bib", "csv", "json"],
                key="lit_file_upload",
                help="支持BibTeX (.bib)、CSV (.csv) 和 JSON (.json) 格式的文献引用文件。",
            )
            if lit_file is not None:
                import tempfile
                with tempfile.NamedTemporaryFile(delete=False, suffix=lit_file.name) as tmp:
                    tmp.write(lit_file.getvalue())
                    tmp_path = tmp.name
                try:
                    engine.import_literature_from_file(tmp_path)
                    st.success(f"已导入文献！")
                    st.rerun()
                except Exception as e:
                    st.error(f"导入失败：{e}")

        st.subheader("文献搜索与管理")

        col_l1, col_l2 = st.columns([3, 1])
        with col_l1:
            lit_search_kw = st.text_input(
                "搜索关键词（用空格分隔）",
                placeholder="例如：自尊 社交焦虑 中介效应",
                key="lit_search_kw",
            )
        with col_l2:
            st.markdown("<br>", unsafe_allow_html=True)
            search_btn = st.button("🔍 搜索文献", width="stretch", key="search_lit_btn")

        # ── 异步文献搜索 pending 状态 ──
        lit_pending = st.session_state.get("_lit_search_pending")
        if lit_pending is not None:
            st.info("⏳ 正在搜索相关文献（含在线数据库），请稍候...")
            col_cancel, _ = st.columns([1, 3])
            with col_cancel:
                if st.button("❌ 取消搜索", width="stretch", key="cancel_lit_search"):
                    cancel_search_request(lit_pending["cancel_id"])
                    try:
                        lit_pending["future"].cancel()
                    except Exception:
                        pass
                    st.session_state.pop("_lit_search_pending", None)
                    st.warning("已取消文献搜索。")
                    st.rerun()

            future = lit_pending["future"]
            if future.done():
                st.session_state.pop("_lit_search_pending", None)
                try:
                    results = future.result()
                    engine.state.suggested_literature = results
                    engine.state.current_step = "literature_searched"
                    engine.state.logs.append(f"✅ 已搜索到{len(results)}篇相关文献")
                    if results:
                        st.success(f"找到 {len(results)} 篇相关文献")
                    else:
                        st.info("未找到相关文献，请尝试其他关键词。")
                except CancelledSearchError:
                    st.warning("文献搜索已被取消。")
                except Exception as e:
                    st.error(f"文献搜索失败：{e}")
                st.rerun()

        if search_btn and lit_search_kw.strip():
            keywords = lit_search_kw.strip().split()
            topic = engine.state.topic or ""
            async_result = search_literature_async(
                keywords=keywords,
                topic=topic,
                include_online=True,
            )
            st.session_state._lit_search_pending = {
                "future": async_result["future"],
                "cancel_id": async_result["cancel_id"],
            }
            st.rerun()

        # 显示文献列表
        all_lit = engine.state.suggested_literature + engine.state.user_literature
        if all_lit:
            st.markdown("**已有文献：**")
            lit_data = []
            for lit in all_lit:
                lit_data.append({
                    "引用键": lit.get("key", ""),
                    "作者": ", ".join(lit.get("authors", [])) if isinstance(lit.get("authors"), list) else lit.get("authors", ""),
                    "年份": lit.get("year", ""),
                    "标题": lit.get("title", ""),
                    "期刊": lit.get("journal", ""),
                    "相关性": lit.get("relevance", ""),
                })
            st.dataframe(pd.DataFrame(lit_data), width="stretch")

        # 手动添加文献
        st.divider()
        st.markdown("**手动添加文献：**")
        col_ua1, col_ua2 = st.columns(2)
        with col_ua1:
            ua_authors = st.text_input("作者（用逗号分隔）", key="ua_authors", placeholder="温忠麟, 叶宝娟")
            ua_year = st.text_input("年份", key="ua_year", placeholder="2014")
            ua_title = st.text_input("标题", key="ua_title", placeholder="中介效应分析: 方法和模型发展")
        with col_ua2:
            ua_journal = st.text_input("期刊", key="ua_journal", placeholder="心理科学进展")
            ua_relevance = st.text_input("与本文的相关性", key="ua_relevance", placeholder="中介分析方法学参考文献")
        ua_is_chinese = st.checkbox("中文文献", value=True, key="ua_is_chinese")

        if st.button("➕ 添加文献", key="add_user_lit"):
            if ua_authors.strip() and ua_title.strip():
                authors_list = [a.strip() for a in ua_authors.split(",") if a.strip()]
                engine.add_user_literature([{
                    "key": f"user_{authors_list[0] if authors_list else 'anon'}{ua_year}",
                    "authors": authors_list,
                    "year": ua_year.strip(),
                    "title": ua_title.strip(),
                    "journal": ua_journal.strip(),
                    "is_chinese": ua_is_chinese,
                    "source": "user",
                    "relevance": ua_relevance.strip(),
                }])
                st.success("文献已添加！")
                st.rerun()
            else:
                st.error("作者和标题为必填项！")

    # ============ 标签页3: 确认问题 ============
    with tab3:
        st.subheader("回答确认问题")
        st.caption("系统将根据您已填写的信息生成确认问题，回答这些问题有助于提高论文质量。")

        if st.button("🔮 生成确认问题", type="primary", key="gen_qa"):
            with st.spinner("正在分析信息缺口并生成问题..."):
                questions = engine.generate_questions()
            st.success(f"已生成 {len(questions)} 个确认问题")
            st.rerun()

        qa_questions = engine.state.qa_questions
        if qa_questions:
            st.markdown(f"共 **{len(qa_questions)}** 个问题需要确认")

            answers = {}
            for q in qa_questions:
                qid = q["id"]
                importance_icon = {"high": "🔴", "medium": "🟡", "low": "🟢"}.get(q.get("importance", "medium"), "⚪")

                with st.expander(f"{importance_icon} [{q.get('category', '')}] {q['question'][:60]}...", expanded=(q.get("importance") == "high")):
                    current_answer = engine.state.qa_answers.get(qid, q.get("default_answer", ""))
                    answer = st.text_area(
                        "您的回答：",
                        value=current_answer,
                        key=f"qa_{qid}",
                        height=80,
                        placeholder=q.get("hint", ""),
                    )
                    if answer and answer != current_answer:
                        answers[qid] = answer

            if answers:
                if st.button("💾 保存回答", type="primary", key="save_qa"):
                    engine.answer_questions(answers)
                    st.success(f"已保存 {len(answers)} 条回答！")
                    st.rerun()

            # 显示未回答的高优问题
            pending = engine.get_pending_gaps()
            if pending:
                st.warning(f"⚠ 还有 {len(pending)} 个高优先级问题未回答，建议在生成论文前完成。")
        else:
            st.info("👆 点击上方按钮生成确认问题。")

    # ============ 标签页4: 生成论文 ============
    with tab4:
        st.subheader("生成论文初稿")

        # 检查信息完整性
        gaps_text = engine.get_gaps_summary()
        if "❌" in gaps_text or "⚠" in gaps_text:
            st.warning("以下信息尚未完善，建议补充后再生成论文：\n\n" + gaps_text)

        if st.button("🚀 生成完整论文初稿", type="primary", width="stretch", key="gen_paper"):
            with st.spinner("正在生成论文初稿（可能需要几十秒）..."):
                try:
                    sections = engine.generate_full_paper()
                    st.success("论文初稿生成完毕！")
                    st.balloons()
                except Exception as e:
                    st.error(f"生成失败：{e}")

        sections = engine.state.generated_sections
        if sections:
            st.divider()
            st.subheader("📄 论文预览")

            # 逐章节显示
            for key, label in section_order:
                content = sections.get(key, "")
                if content:
                    with st.expander(f"{label}", expanded=(key in ["title", "abstract"])):
                        if key == "title":
                            st.markdown(f"# {content}")
                        elif key == "abstract":
                            st.markdown(f"**摘要：** {content}")
                        elif key == "keywords":
                            st.markdown(f"**关键词：** {content}")
                        elif key == "references":
                            # references is a list
                            if isinstance(content, list):
                                for ref in content:
                                    st.markdown(ref)
                            else:
                                st.markdown(content)
                        else:
                            st.markdown(content)

            # 导出
            st.divider()
            st.subheader("📥 导出论文")

            export_fmt = st.radio(
                "选择导出格式",
                ["📄 Markdown (.md)", "📕 Word (.docx)", "🌐 HTML (.html)"],
                horizontal=True,
                key="paper_export_fmt",
            )

            if st.button("📥 导出论文", type="primary", key="export_paper"):
                manuscript = engine.assemble_manuscript()
                if export_fmt.startswith("📄 Markdown"):
                    import base64
                    b64 = base64.b64encode(manuscript.encode("utf-8")).decode()
                    title = sections.get("title", "论文初稿").replace("/", "-")
                    href = f'<a href="data:text/markdown;base64,{b64}" download="{title}.md">点击下载 Markdown 文件 (.md)</a>'
                    st.markdown(href, unsafe_allow_html=True)
                    st.success("✅ Markdown 文件已生成。")

                elif export_fmt.startswith("📕 Word"):
                    try:
                        import base64
                        from docx import Document
                        from docx.shared import Pt, Cm
                        import io

                        doc = Document()
                        style = doc.styles["Normal"]
                        font = style.font
                        font.name = "SimSun"
                        font.size = Pt(12)

                        title_text = sections.get("title", "心理学实证研究")
                        doc.add_heading(title_text, level=0)

                        doc.add_heading("摘要", level=1)
                        doc.add_paragraph(sections.get("abstract", ""))
                        doc.add_paragraph(f"关键词：{sections.get('keywords', '')}")

                        for key, label in section_order:
                            if key in ["title", "abstract", "keywords", "references"]:
                                continue
                            content = sections.get(key, "")
                            if content:
                                doc.add_heading(label, level=1)
                                for para in content.split("\n\n"):
                                    if para.strip():
                                        doc.add_paragraph(para.strip())

                        doc.add_heading("参考文献", level=1)
                        refs = sections.get("references", [])
                        if isinstance(refs, list):
                            for ref in refs:
                                doc.add_paragraph(ref, style="List Number")
                        else:
                            doc.add_paragraph(refs)

                        buf = io.BytesIO()
                        doc.save(buf)
                        b64 = base64.b64encode(buf.getvalue()).decode()
                        title = title_text.replace("/", "-")
                        href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{title}.docx">点击下载 Word 文档 (.docx)</a>'
                        st.markdown(href, unsafe_allow_html=True)
                        st.success("✅ Word 文档已生成。")
                    except Exception as e:
                        st.error(f"Word 导出失败：{e}。已为您准备 Markdown 版本。")
                        import base64
                        b64 = base64.b64encode(manuscript.encode("utf-8")).decode()
                        href = f'<a href="data:text/markdown;base64,{b64}" download="论文初稿.md">点击下载 Markdown 文件 (.md)</a>'
                        st.markdown(href, unsafe_allow_html=True)

                elif export_fmt.startswith("🌐 HTML"):
                    import base64
                    title = sections.get("title", "论文初稿")
                    html_content = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<title>{title}</title>
<style>
body {{ font-family: 'SimSun', 'Microsoft YaHei', serif; max-width: 800px; margin: 0 auto; padding: 2rem; line-height: 2; }}
h1 {{ text-align: center; font-size: 1.5em; }}
h2 {{ border-bottom: 1px solid #333; }}
ref {{ font-size: 0.9em; }}
</style>
</head>
<body>
{manuscript.replace(chr(10), '<br>').replace('# ', '<h1>').replace('## ', '<h2>').replace('---', '<hr>').replace('**', '<b>').replace('</b><b>', '')}
</body>
</html>"""
                    b64 = base64.b64encode(html_content.encode("utf-8")).decode()
                    title_safe = title.replace("/", "-")
                    href = f'<a href="data:text/html;base64,{b64}" download="{title_safe}.html">点击下载 HTML 文件 (.html)</a>'
                    st.markdown(href, unsafe_allow_html=True)
                    st.success("✅ HTML 文件已生成。")

    # ============ 标签页5: 答辩模拟 ============
    with tab5:
        _render_paper_mode_defense_qa()

    # 如果没有分析结果但有分析输出，提示导入
    if st.session_state.analysis_output is not None and not engine.state.analysis_results:
        with tab1:
            with st.expander("📥 从数据分析导入结果", expanded=False):
                st.info("检测到数据分析输出，可以导入到论文写作中。")
                if st.button("📥 导入当前分析结果", key="import_analysis"):
                    output = st.session_state.analysis_output
                    test_type = output.get("test_type", "unknown")
                    engine.import_analysis_results({test_type: output})
                    st.success("分析结果已导入！")
                    st.rerun()

    st.divider()
    st.caption("💡 提示：生成的论文为初稿，请结合专业知识进行审阅和修改。参考文献请在使用前核实其准确性。")


def _render_paper_mode_defense_qa():
    """v4.6 F3: 论文写作 mode 内嵌答辩模拟。

    依赖 st.session_state.plan + st.session_state.analysis_output；
    若两者缺失则提示用户先去「📈 数据分析」mode 跑分析。
    """
    st.subheader("🎤 答辩问题预演")
    st.caption("根据你刚跑完的数据分析，自动生成老师答辩时最可能问的问题与模板答案。")

    plan = st.session_state.get("plan")
    output = st.session_state.get("analysis_output")
    if plan is None or output is None:
        st.info(
            "尚未检测到分析结果。请先到「📈 数据分析」mode 完成一次分析，"
            "再回来生成针对性的答辩问题。"
        )
        return

    test_type = output.get("test_type", "")
    test_name = output.get("test_name_zh", test_type or "未知方法")
    st.caption(f"当前分析方法：**{test_name}**")

    max_items = st.slider(
        "生成问题数量", min_value=3, max_value=10, value=7,
        key="paper_defense_qa_max",
    )

    if st.button("🎤 生成答辩问题", type="secondary",
                 width="stretch", key="paper_defense_qa_gen"):
        ctx = {"test_type": test_type, "test_name_zh": test_name}
        with st.spinner("正在生成针对性答辩问题..."):
            items = generate_defense_qa(plan=plan, output=output, ctx=ctx,
                                         max_items=max_items)
        st.session_state["_paper_defense_qa_items"] = items

    items = st.session_state.get("_paper_defense_qa_items", [])
    if not items:
        return

    st.divider()
    st.info(
        "**难度分级：** 🟢 必问（务必准备） · 🟡 常问（建议准备） · 🔴 刁钻（视情况）"
    )
    st.markdown(f"**已生成 {len(items)} 个问题，按难度排序：**")

    groups = group_qa_by_difficulty(items)
    counter = 0
    mastered_map = st.session_state.setdefault("paper_defense_qa_mastered", {})
    for diff in ("必问", "常问", "刁钻"):
        diff_items = groups.get(diff, [])
        if not diff_items:
            continue
        emoji = {"必问": "🟢", "常问": "🟡", "刁钻": "🔴"}[diff]
        st.markdown(f"#### {emoji} {diff}")
        for item in diff_items:
            counter += 1
            with st.container():
                cols_q = st.columns([10, 2])
                cols_q[0].markdown(f"**Q{counter}：{item.question}**")
                new_mastered = cols_q[1].checkbox(
                    "✅ 已掌握",
                    value=mastered_map.get(item.question_id, False),
                    key=f"paper_qa_mastered_{item.question_id}",
                )
                mastered_map[item.question_id] = new_mastered
                st.caption(f"{item.category_label} · {emoji} {diff}")
                st.markdown(f"💬 **答**：{item.answer}")
                st.markdown("")

    st.divider()
    md = render_qa_as_markdown(items)
    with st.expander("📋 复制全部问答（Markdown 格式）", expanded=False):
        st.code(md, language="markdown")
