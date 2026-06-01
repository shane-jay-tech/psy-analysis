"""问卷设计报告导出：Word (DOCX) / PDF"""

import io
import os
from typing import Dict, List, Optional


# ================================================================
# Word (DOCX) 导出
# ================================================================


def export_to_docx(design: Dict) -> bytes:
    """生成格式化的 Word 文档，返回 bytes。"""
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()

    # --- 页面设置 ---
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # --- 样式 ---
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.8
    _set_east_asian_font(style, "宋体")

    for level, (size, bold) in {1: (22, True), 2: (16, True), 3: (14, True)}.items():
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.name = "黑体"
        h_style.font.size = Pt(size)
        h_style.font.bold = bold
        h_style.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        _set_east_asian_font(h_style, "黑体")

    # --- 封面 ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(60)
    run = title_para.add_run("心理学问卷设计报告")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.name = "黑体"
    _set_east_asian_font(run, "黑体")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"—— {design['construct_name']}问卷")
    run.font.size = Pt(16)
    run.font.name = "楷体"

    doc.add_paragraph()

    # --- 研究问题 ---
    doc.add_heading("研究问题", level=2)
    p = doc.add_paragraph()
    run = p.add_run(design["research_question"])
    run.font.size = Pt(12)
    run.bold = True

    # --- 构念识别 ---
    doc.add_heading("构念识别与设计思路", level=2)

    sc = design["scale_config"]
    template = design["template_used"]

    info_lines = [
        f"识别构念：{design['construct_name']}",
        f"识别方式：{'大语言模型智能分析' if design.get('llm_used') else '内置构念知识库匹配'}",
        f"题型：{template['name']}（{sc['points']}点 Likert 量表）",
        f"总题量：{sc['n_items']} 题（{sc['n_dimensions']} 个维度）",
        f"反向题：{sc['n_reverse']} 题（占比 {sc['reverse_ratio']}）",
        f"预计用时：约 {max(3, sc['n_items'] // 3)} 分钟",
    ]
    for line in info_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.first_line_indent = Cm(0.74)

    doc.add_paragraph(design["match_reason"])

    # --- 构念定义 ---
    doc.add_heading("构念定义与理论框架", level=2)
    construct = design.get("matched_construct") or {}
    definition = (
        construct.get("definition")
        or design.get("llm_definition")
        or f"{design['construct_name']}的操作性定义将基于文献综述确定。"
    )
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.add_run(definition)

    # --- 维度结构 ---
    doc.add_heading("维度结构", level=2)
    doc.add_paragraph(f"本问卷将「{design['construct_name']}」分解为 {len(design['dimensions_used'])} 个理论维度：")

    for i, dim in enumerate(design["dimensions_used"]):
        doc.add_heading(f"维度 {i+1}：{dim['name']}", level=3)
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.add_run(f"描述：{dim['desc']}")
        doc.add_paragraph(f"题目数：{dim.get('item_count', '-')} 题")
        if dim.get("example"):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)
            run = p.add_run('示例：“' + dim['example'] + '”')
            run.italic = True

    # --- 量表技术参数 ---
    doc.add_heading("量表技术参数", level=2)
    table = doc.add_table(rows=7, cols=2, style="Light Grid Accent 1")
    params = [
        ("参数", "设置"),
        ("题型", template["name"]),
        ("量表点数", f"{sc['points']} 点 Likert"),
        ("总题量", f"{sc['n_items']} 题"),
        ("维度数", str(sc["n_dimensions"])),
        ("反向题数", f"{sc['n_reverse']} 题（{sc['reverse_ratio']}）"),
        ("预计用时", f"约 {max(3, sc['n_items'] // 3)} 分钟"),
    ]
    for row_idx, (label, value) in enumerate(params):
        table.cell(row_idx, 0).text = label
        table.cell(row_idx, 1).text = value
        if row_idx == 0:
            for cell in table.rows[row_idx].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    # --- 评分锚定 ---
    doc.add_heading("评分锚定", level=3)
    for anchor in sc.get("anchors", []):
        doc.add_paragraph(anchor, style="List Bullet")

    # --- 指导语 ---
    doc.add_heading("问卷指导语", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.add_run(design["instructions"])

    # --- 问卷题目 ---
    doc.add_heading("问卷题目", level=2)

    current_dim = None
    for item in design["items"]:
        if item["dimension"] != current_dim:
            current_dim = item["dimension"]
            doc.add_heading(f"▎ {current_dim}", level=3)
        rev_mark = " 【反向计分】" if item["reverse"] else ""
        text = f"Q{item['index']}. {item['text']}{rev_mark}"
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(0.74)
        doc.add_paragraph("[1]    [2]    [3]    [4]    [5]")

    # --- 计分方式 ---
    doc.add_heading("计分方式", level=2)
    doc.add_paragraph(design["scoring"])

    rev_items = [it for it in design["items"] if it["reverse"]]
    if rev_items:
        p = doc.add_paragraph("需反向计分的题目：")
        p.add_run(", ".join(f"Q{it['index']}" for it in rev_items)).bold = True

    # --- 信效度保障 ---
    doc.add_heading("信效度保障策略", level=2)
    for section_name, content in design["psychometrics"].items():
        doc.add_heading(section_name, level=3)
        for line in content.split("\n"):
            stripped = line.strip()
            if stripped:
                doc.add_paragraph(stripped, style="List Bullet")

    # --- 施测建议 ---
    doc.add_heading("施测建议", level=2)
    suggestions = [
        f"施测对象：根据研究问题确定的目标群体",
        f"施测方式：纸笔或在线问卷（推荐使用问卷星/Qualtrics等平台）",
        f"施测时间：约 {max(3, sc['n_items'] // 3)}-{max(5, sc['n_items'] // 2)} 分钟",
        f"预测试：正式施测前，建议选取30-50名目标被试进行预测试，检验题目理解度和信度",
        f"正式施测样本量：建议 N ≥ {sc['n_items'] * 10}（基于EFA的样本量要求）",
    ]
    for s in suggestions:
        doc.add_paragraph(s, style="List Bullet")

    # --- 参考文献 ---
    doc.add_heading("参考文献", level=2)

    if construct and construct.get("references"):
        doc.add_heading("构念相关文献", level=3)
        for i, ref in enumerate(construct["references"]):
            doc.add_paragraph(f"{i+1}. {ref}")

    if design.get("llm_references"):
        doc.add_heading("LLM 生成的参考文献（请核实后再引用）", level=3)
        for i, ref in enumerate(design["llm_references"]):
            doc.add_paragraph(f"{i+1}. {ref}")

    doc.add_heading("测量学通用参考文献", level=3)
    general_refs = [
        "DeVellis, R. F., & Thorpe, C. T. (2021). Scale Development: Theory and Applications (5th ed.). SAGE.",
        "Nunnally, J. C., & Bernstein, I. H. (1994). Psychometric Theory (3rd ed.). McGraw-Hill.",
        "Furr, R. M. (2017). Psychometrics: An Introduction (3rd ed.). SAGE.",
        "Haynes, S. N., Richard, D. C. S., & Kubany, E. S. (1995). Content validity in psychological assessment: A functional approach to concepts and methods. Psychological Assessment, 7(3), 238-247.",
        "Hinkin, T. R. (1998). A brief tutorial on the development of measures for use in survey questionnaires. Organizational Research Methods, 1(1), 104-121.",
        "Hu, L., & Bentler, P. M. (1999). Cutoff criteria for fit indexes in covariance structure analysis. Structural Equation Modeling, 6(1), 1-55.",
    ]
    for i, ref in enumerate(general_refs):
        doc.add_paragraph(f"{i+1}. {ref}")

    # --- 已有成熟量表 ---
    established = []
    if construct and construct.get("established_scales"):
        established.extend(construct["established_scales"])
    if design.get("llm_established_scales"):
        established.extend(design["llm_established_scales"])

    if established:
        doc.add_heading("已有成熟量表参考", level=2)
        for scale in established:
            doc.add_paragraph(scale, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _set_east_asian_font(style, font_name: str) -> None:
    """设置样式的东亚字体（CJK fallback）。不是关键功能，失败时静默回退。"""
    try:
        from docx.oxml.ns import qn
        from lxml import etree
        rPr = style.element.find(qn("w:rPr"))
        if rPr is None:
            rPr = etree.SubElement(style.element, qn("w:rPr"))
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = etree.SubElement(rPr, qn("w:rFonts"))
        rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass


# ================================================================
# PDF 导出
# ================================================================


def export_to_pdf(design: Dict) -> bytes:
    """生成 PDF 文件，使用 fpdf2 并嵌入中文字体，返回 bytes。"""
    from fpdf import FPDF
    from src.visualization.fonts import find_chinese_font

    font_path = find_chinese_font()
    if not font_path:
        raise RuntimeError("未找到可用的中文字体文件，无法生成 PDF。")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # 注册中文字体
    font_name = "CJK"
    pdf.add_font(font_name, "", font_path, uni=True)
    pdf.add_font(font_name, "B", font_path, uni=True)

    def write(text, size=12, bold=False, align="L", new_line=True):
        style = "B" if bold else ""
        pdf.set_font(font_name, style, size)
        width = pdf.epw  # effective page width
        if new_line:
            pdf.cell(width, 8, text, new_x="LMARGIN", new_y="NEXT", align=align)
        else:
            pdf.cell(width, 8, text, align=align)

    def write_multi(text, size=12, bold=False, align="L"):
        style = "B" if bold else ""
        pdf.set_font(font_name, style, size)
        pdf.multi_cell(pdf.epw, 7, text, align=align)

    # --- 封面 ---
    pdf.ln(20)
    write("心理学问卷设计报告", size=22, bold=True, align="C")
    write(f"—— {design['construct_name']}问卷", size=14, align="C")
    pdf.ln(8)

    # --- 研究问题 ---
    pdf.ln(4)
    write("一、研究问题", size=16, bold=True)
    write_multi(design["research_question"], size=12, bold=True)

    # --- 构念识别 ---
    pdf.ln(4)
    write("二、构念识别与设计思路", size=16, bold=True)
    sc = design["scale_config"]
    template = design["template_used"]

    info_lines = [
        f"识别构念：{design['construct_name']}",
        f"识别方式：{'大语言模型智能分析' if design.get('llm_used') else '内置构念知识库匹配'}",
        f"题型：{template['name']}（{sc['points']} 点 Likert 量表）",
        f"总题量：{sc['n_items']} 题（{sc['n_dimensions']} 个维度），反向题 {sc['n_reverse']} 题",
        f"预计用时：约 {max(3, sc['n_items'] // 3)} 分钟",
    ]
    for line in info_lines:
        write_multi(line, size=11)

    write_multi(design["match_reason"], size=11)

    # --- 构念定义 ---
    pdf.ln(4)
    write("三、构念定义与理论框架", size=16, bold=True)
    construct = design.get("matched_construct") or {}
    definition = (
        construct.get("definition")
        or design.get("llm_definition")
        or f"{design['construct_name']}的操作性定义将基于文献综述确定。"
    )
    write_multi(definition, size=11)

    # --- 维度结构 ---
    pdf.ln(4)
    write("四、维度结构", size=16, bold=True)
    write_multi(f"本问卷包含 {len(design['dimensions_used'])} 个理论维度：", size=11)

    for i, dim in enumerate(design["dimensions_used"]):
        write(f"{i+1}. {dim['name']}：{dim['desc']}（{dim.get('item_count', '?')}题）", size=11, bold=True)
        if dim.get("example"):
            write_multi('   示例："' + dim['example'] + '"', size=11)

    # --- 技术参数 ---
    pdf.ln(4)
    write("五、量表技术参数", size=16, bold=True)
    pdf.set_font(font_name, "", 11)
    col_w = [50, 120]
    params = [
        ("题型", template["name"]),
        ("量表点数", f"{sc['points']} 点"),
        ("总题量", f"{sc['n_items']} 题"),
        ("维度数", str(sc['n_dimensions'])),
        ("反向题", f"{sc['n_reverse']} 题 ({sc['reverse_ratio']})"),
    ]
    for label, value in params:
        pdf.cell(col_w[0], 7, label + "：", new_x="RIGHT", new_y="LAST")
        pdf.cell(col_w[1], 7, value, new_x="LMARGIN", new_y="NEXT")

    pdf.ln(2)
    write("评分锚定：", size=11, bold=True)
    for anchor in sc.get("anchors", []):
        write_multi(f"  {anchor}", size=11)

    # --- 指导语 ---
    pdf.ln(4)
    write("六、问卷指导语", size=16, bold=True)
    pdf.set_left_margin(25)
    write_multi(design["instructions"], size=11)
    pdf.set_left_margin(20)

    # --- 问卷题目 ---
    pdf.ln(4)
    write("七、问卷题目", size=16, bold=True)

    current_dim = None
    for item in design["items"]:
        if item["dimension"] != current_dim:
            current_dim = item["dimension"]
            pdf.ln(2)
            write(f"▎ {current_dim}", size=13, bold=True)
        rev_mark = " 【反向计分】" if item["reverse"] else ""
        write_multi(f"Q{item['index']}. {item['text']}{rev_mark}", size=11)
        write("    [1]  [2]  [3]  [4]  [5]", size=10)

    # --- 计分 ---
    pdf.ln(4)
    write("八、计分方式", size=16, bold=True)
    write_multi(design["scoring"], size=11)

    rev_items = [it for it in design["items"] if it["reverse"]]
    if rev_items:
        write("需反向计分的题目：", size=11, bold=True)
        write_multi(", ".join(f"Q{it['index']}" for it in rev_items), size=11)

    # --- 信效度 ---
    pdf.ln(4)
    write("九、信效度保障策略", size=16, bold=True)
    for section_name, content in design["psychometrics"].items():
        pdf.ln(1)
        write(f"{section_name}", size=13, bold=True)
        for line in content.split("\n"):
            stripped = line.strip()
            if stripped:
                write_multi(f"  • {stripped}", size=10)

    # --- 施测建议 ---
    pdf.ln(4)
    write("十、施测建议", size=16, bold=True)
    suggestions = [
        f"施测对象：根据研究问题确定的目标群体",
        f"施测方式：纸笔或在线问卷",
        f"施测时间：约 {max(3, sc['n_items'] // 3)}-{max(5, sc['n_items'] // 2)} 分钟",
        f"预测试：30-50名目标被试",
        f"样本量：N ≥ {sc['n_items'] * 10}（EFA要求）",
    ]
    for s in suggestions:
        write_multi(f"  • {s}", size=11)

    # --- 参考文献 ---
    pdf.ln(4)
    write("十一、参考文献", size=16, bold=True)

    all_refs = []
    if construct and construct.get("references"):
        all_refs.extend(construct["references"])
    if design.get("llm_references"):
        all_refs.extend(design["llm_references"])
    all_refs.extend([
        "DeVellis, R. F., & Thorpe, C. T. (2021). Scale Development (5th ed.). SAGE.",
        "Nunnally, J. C., & Bernstein, I. H. (1994). Psychometric Theory (3rd ed.). McGraw-Hill.",
        "Hinkin, T. R. (1998). A brief tutorial on the development of measures. Organizational Research Methods, 1(1), 104-121.",
    ])

    for i, ref in enumerate(all_refs):
        write_multi(f"{i+1}. {ref}", size=10)

    return pdf.output()


# ================================================================
# v4.1：正式问卷 PDF 导出（输入 = ItemsDoc）
# ================================================================


_DEFAULT_LIKERT_ANCHORS_ZH = {
    5: ["完全不符合", "比较不符合", "一般", "比较符合", "完全符合"],
    7: ["完全不符合", "不符合", "比较不符合", "一般", "比较符合", "符合", "完全符合"],
}


def build_questionnaire_pdf(
    items_doc,
    *,
    scale_points: int = 5,
    anchors: Optional[List[str]] = None,
    header_meta: Optional[Dict[str, str]] = None,
    show_id_field: bool = True,
) -> bytes:
    """生成正式问卷 PDF（v4.1）。

    输入 ItemsDoc（src.questionnaire.items_loader），输出 .pdf bytes。
    与 build_questionnaire_docx 同一布局：标题 + 指导语 + 锚点 + 题目（每题下显示 [1]..[N] 选项）+ 致谢。
    """
    from fpdf import FPDF
    from src.visualization.fonts import find_chinese_font

    if scale_points < 2 or scale_points > 11:
        raise ValueError(f"scale_points 需在 2-11 之间（传入 {scale_points}）")
    if anchors is None:
        anchors = _DEFAULT_LIKERT_ANCHORS_ZH.get(
            scale_points,
            [str(i + 1) for i in range(scale_points)],
        )
    if len(anchors) != scale_points:
        raise ValueError(
            f"anchors 长度 ({len(anchors)}) 需与 scale_points ({scale_points}) 一致"
        )

    title = (getattr(items_doc, "title", "") or "心理测量问卷").strip()
    instructions = (getattr(items_doc, "instructions", "") or "").strip()
    items = list(getattr(items_doc, "items", []) or [])
    reverse_set = set(getattr(items_doc, "reverse_indices", []) or [])
    if not items:
        raise ValueError("ItemsDoc.items 为空，无法生成问卷文档。")

    font_path = find_chinese_font()
    if not font_path:
        raise RuntimeError("未找到可用的中文字体文件，无法生成 PDF。")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    font_name = "CJK"
    pdf.add_font(font_name, "", font_path)
    pdf.add_font(font_name, "B", font_path)

    def write(text: str, size=12, bold=False, align="L", new_line=True):
        style = "B" if bold else ""
        pdf.set_font(font_name, style, size)
        if new_line:
            pdf.cell(pdf.epw, 8, text, new_x="LMARGIN", new_y="NEXT", align=align)
        else:
            pdf.cell(pdf.epw, 8, text, align=align)

    def write_multi(text: str, size=11, bold=False, align="L"):
        style = "B" if bold else ""
        pdf.set_font(font_name, style, size)
        pdf.multi_cell(pdf.epw, 7, text, align=align)

    # 标题
    pdf.ln(8)
    write(title, size=20, bold=True, align="C")
    pdf.ln(2)

    # 元信息
    if header_meta:
        meta_parts = []
        for key, label in [
            ("researcher", "主试"),
            ("project", "研究项目"),
            ("version", "版本"),
            ("date", "日期"),
        ]:
            val = header_meta.get(key)
            if val:
                meta_parts.append(f"{label}：{val}")
        if meta_parts:
            write("    ".join(meta_parts), size=10, align="C")

    # 编号填写区
    if show_id_field:
        pdf.ln(2)
        write("编号：__________    性别：____    年龄：____    日期：__________",
              size=11, align="L")

    # 指导语
    if instructions:
        pdf.ln(4)
        write("指导语", size=14, bold=True)
        write_multi(instructions, size=11)

    # 评分说明
    pdf.ln(4)
    write("评分说明", size=14, bold=True)
    anchor_text = "；".join(f"{i + 1} = {anchors[i]}" for i in range(scale_points))
    write_multi(
        "请在每道题后选择最符合您实际情况的数字（圈选）：" + anchor_text + "。",
        size=11,
    )
    if reverse_set:
        write_multi(
            f"题号后标 (R) 的为反向题，共 {len(reverse_set)} 题，计分时需反向。",
            size=11,
        )

    # 题项
    pdf.ln(4)
    write("题项", size=14, bold=True)
    score_row = "    ".join(f"[{i + 1}]" for i in range(scale_points))
    for idx, item in enumerate(items):
        is_rev = idx in reverse_set
        rev_mark = " (R)" if is_rev else ""
        write_multi(f"{idx + 1}.{rev_mark}  {item}", size=11)
        write_multi("    " + score_row, size=10)
        pdf.ln(1)

    # 致谢
    pdf.ln(4)
    write("—— 问卷结束，感谢您的参与 ——", size=12, bold=True, align="C")

    out = pdf.output()
    if isinstance(out, (bytes, bytearray)):
        return bytes(out)
    return bytes(out)  # fpdf2 新版本返回 bytearray，统一转 bytes
