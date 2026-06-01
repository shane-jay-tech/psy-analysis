"""APA7 中文论文 Word 导出器。

输入：研究元信息 + 方法/结果 Markdown 段落 + 描述统计表 + 图表 PNG bytes
输出：.docx 字节流（可由 Streamlit st.download_button 直接发给浏览器）

设计原则：
- 不依赖外部模板，纯代码生成（避免模板路径问题）
- 中英文混排自动设字体（latin/east-asia 双轨）
- 简易 Markdown 解析：识别 # / ## / ### 标题、空行分段、**bold**、*italic*
  完整 Markdown 解析交给后续版本（v2.7 仅覆盖向导生成的简单格式）
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import List, Optional, Tuple

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from . import docx_styles as S


@dataclass
class ThesisMeta:
    """论文元信息。"""

    title: str = "心理学实证研究报告"
    author: str = ""
    affiliation: str = ""
    date: str = ""
    abstract: str = ""
    keywords: List[str] = field(default_factory=list)


@dataclass
class FigureItem:
    """图表条目。"""

    caption: str
    png_bytes: bytes
    width_cm: float = 12.0


def plotly_figs_to_figure_items(
    plotly_specs: List[Tuple[str, "object"]],
    palette: str = "grayscale",
    *,
    width_px: int = 1500,
    height_px: int = 1000,
    width_cm: float = 12.0,
) -> List[FigureItem]:
    """便捷函数：把 [(caption, plotly_fig)] 转成 FigureItem 列表。

    内部统一调用 paper_export.to_paper_png()，保证与"下载论文版图表"按钮
    完全一致的输出（300dpi 等价 + 学术配色）。

    Args:
        plotly_specs: 列表 of (中文图注, plotly Figure)
        palette: "grayscale" / "color" / "mono"，与 paper_export 一致
        width_px / height_px: PNG 像素尺寸
        width_cm: 嵌入 Word 时的宽度（厘米）

    Returns:
        FigureItem 列表；kaleido 缺失时静默跳过该图，返回部分结果。
    """
    from src.visualization.paper_export import (
        KaleidoMissingError, to_paper_png,
    )
    items: List[FigureItem] = []
    for caption, fig in plotly_specs:
        try:
            png = to_paper_png(
                fig, palette=palette,
                width_px=width_px, height_px=height_px,
            )
            items.append(FigureItem(caption=caption, png_bytes=png, width_cm=width_cm))
        except KaleidoMissingError:
            break  # kaleido 不可用，停止后续生成
        except Exception:
            continue
    return items


# --------------------------------------------------------------------------- #
# 字体设置工具
# --------------------------------------------------------------------------- #

def _set_run_font(run, *, size: float = S.SIZE_BODY, bold: bool = False,
                  italic: bool = False, latin: str = S.FONT_BODY_LATIN,
                  cjk: str = S.FONT_BODY_CJK, color_hex: Optional[str] = None):
    run.font.name = latin
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    # 中文字体（east-asia）必须通过 XML 设置
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), cjk)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)


def _add_paragraph(doc, text: str = "", *, level: int = 0, align: str = "left",
                   first_line_indent: bool = False, italic_runs: Optional[List[bool]] = None):
    """添加段落，level 0=正文, 1=H1, 2=H2, 3=H3。"""
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    pf = p.paragraph_format
    pf.line_spacing = S.LINE_SPACING
    if level == 0 and first_line_indent:
        pf.first_line_indent = Cm(S.FIRST_LINE_INDENT_CM)
    if level > 0:
        pf.space_before = Pt(S.SPACE_BEFORE_HEADING_PT)
        pf.space_after = Pt(S.SPACE_AFTER_HEADING_PT)

    size_map = {0: S.SIZE_BODY, 1: S.SIZE_H1, 2: S.SIZE_H2, 3: S.SIZE_H3}
    size = size_map.get(level, S.SIZE_BODY)
    bold = level > 0
    cjk_font = S.FONT_HEADING_CJK if level > 0 else S.FONT_BODY_CJK
    latin_font = S.FONT_HEADING_LATIN if level > 0 else S.FONT_BODY_LATIN

    if text:
        run = p.add_run(text)
        _set_run_font(run, size=size, bold=bold, latin=latin_font, cjk=cjk_font)
    return p


_INLINE_ITALIC = re.compile(r"\*([^*\n]+)\*")
_INLINE_BOLD = re.compile(r"\*\*([^*\n]+)\*\*")


def _add_paragraph_with_inline_format(doc, text: str, *, first_line_indent: bool = True):
    """正文段落，识别 **bold** 和 *italic*。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = S.LINE_SPACING
    if first_line_indent:
        pf.first_line_indent = Cm(S.FIRST_LINE_INDENT_CM)

    # 简易解析：先按 ** 切，再每段按 * 切
    pos = 0
    for bold_match in _INLINE_BOLD.finditer(text):
        # bold 之前的部分（可能含 italic）
        before = text[pos:bold_match.start()]
        _emit_italic_aware(p, before, bold=False)
        # bold 内容
        _emit_italic_aware(p, bold_match.group(1), bold=True)
        pos = bold_match.end()
    _emit_italic_aware(p, text[pos:], bold=False)
    return p


def _emit_italic_aware(p, text: str, bold: bool):
    if not text:
        return
    pos = 0
    for italic_match in _INLINE_ITALIC.finditer(text):
        before = text[pos:italic_match.start()]
        if before:
            run = p.add_run(before)
            _set_run_font(run, size=S.SIZE_BODY, bold=bold, italic=False)
        run = p.add_run(italic_match.group(1))
        _set_run_font(run, size=S.SIZE_BODY, bold=bold, italic=True)
        pos = italic_match.end()
    rest = text[pos:]
    if rest:
        run = p.add_run(rest)
        _set_run_font(run, size=S.SIZE_BODY, bold=bold, italic=False)


# --------------------------------------------------------------------------- #
# Markdown -> docx 段落
# --------------------------------------------------------------------------- #

def _render_markdown(doc, md: str):
    """把简易 Markdown 转换成 Word 段落。

    支持的语法：
    - # / ## / ### 标题
    - 空行 = 分段
    - **bold** / *italic*
    - - 列表项（简单实现）
    """
    if not md:
        return
    lines = md.split("\n")
    buffer: List[str] = []

    def flush_paragraph():
        if buffer:
            text = " ".join(s.strip() for s in buffer if s.strip())
            if text:
                _add_paragraph_with_inline_format(doc, text, first_line_indent=True)
            buffer.clear()

    for line in lines:
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            continue
        if stripped.startswith("### "):
            flush_paragraph()
            _add_paragraph(doc, stripped[4:], level=3)
        elif stripped.startswith("## "):
            flush_paragraph()
            _add_paragraph(doc, stripped[3:], level=2)
        elif stripped.startswith("# "):
            flush_paragraph()
            _add_paragraph(doc, stripped[2:], level=1)
        elif stripped.startswith("---"):
            flush_paragraph()
            # 分隔线 → 加一个空段
            _add_paragraph(doc, "")
        elif stripped.startswith("- "):
            flush_paragraph()
            p = doc.add_paragraph(style=None)
            p.paragraph_format.line_spacing = S.LINE_SPACING
            p.paragraph_format.left_indent = Cm(0.74)
            run = p.add_run("• " + stripped[2:])
            _set_run_font(run, size=S.SIZE_BODY)
        else:
            buffer.append(stripped)

    flush_paragraph()


# --------------------------------------------------------------------------- #
# 表格
# --------------------------------------------------------------------------- #

def _add_dataframe_table(doc, df: pd.DataFrame, *, caption: str = "", number: int = 1):
    """三线表风格。"""
    if caption:
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cap_p.paragraph_format.space_before = Pt(12)
        cap_p.paragraph_format.space_after = Pt(4)
        run = cap_p.add_run(f"{S.TABLE_PREFIX}{number}  {caption}")
        _set_run_font(run, size=S.SIZE_BODY, bold=True, cjk=S.FONT_HEADING_CJK)

    n_rows = len(df) + 1
    n_cols = len(df.columns)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 表头
    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(str(col))
        _set_run_font(run, size=S.SIZE_TABLE, bold=True, cjk=S.FONT_HEADING_CJK)

    # 数据
    for i, (_, row) in enumerate(df.iterrows(), start=1):
        for j, col in enumerate(df.columns):
            cell = table.rows[i].cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            val = row[col]
            if isinstance(val, float):
                txt = f"{val:.3f}" if abs(val) < 1000 else f"{val:.2f}"
            else:
                txt = str(val) if pd.notna(val) else ""
            run = cell.paragraphs[0].add_run(txt)
            _set_run_font(run, size=S.SIZE_TABLE)

    _apply_three_line_borders(table)


def _apply_three_line_borders(table):
    """三线表：仅顶部、表头底、表尾有横线。"""
    from docx.oxml import OxmlElement

    def set_cell_border(cell, edge: str, present: bool):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.find(qn("w:tcBorders"))
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)
        # remove existing edge
        existing = tc_borders.find(qn(f"w:{edge}"))
        if existing is not None:
            tc_borders.remove(existing)
        border = OxmlElement(f"w:{edge}")
        if present:
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), str(S.TABLE_BORDER_SIZE))
            border.set(qn("w:color"), S.TABLE_BORDER_COLOR)
        else:
            border.set(qn("w:val"), "nil")
        tc_borders.append(border)

    n_rows = len(table.rows)
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell, "top", i == 0)
            set_cell_border(cell, "bottom", i == 0 or i == n_rows - 1)
            set_cell_border(cell, "left", False)
            set_cell_border(cell, "right", False)
            set_cell_border(cell, "insideH", False)
            set_cell_border(cell, "insideV", False)


# --------------------------------------------------------------------------- #
# 图片
# --------------------------------------------------------------------------- #

def _add_figure(doc, fig: FigureItem, number: int = 1):
    """嵌入 PNG 并加图注。"""
    # 图片居中
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(io.BytesIO(fig.png_bytes), width=Cm(fig.width_cm))

    # 图注
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_after = Pt(12)
    cap_run = cap_p.add_run(f"{S.FIG_PREFIX}{number}  {fig.caption}")
    _set_run_font(cap_run, size=S.SIZE_CAPTION, bold=True, cjk=S.FONT_HEADING_CJK)


# --------------------------------------------------------------------------- #
# 主入口
# --------------------------------------------------------------------------- #

def _setup_page(doc):
    """页边距 + 默认样式。"""
    for section in doc.sections:
        section.top_margin = Cm(S.MARGIN_TOP_CM)
        section.bottom_margin = Cm(S.MARGIN_BOTTOM_CM)
        section.left_margin = Cm(S.MARGIN_LEFT_CM)
        section.right_margin = Cm(S.MARGIN_RIGHT_CM)

    # 修改 Normal 样式默认字体
    style = doc.styles["Normal"]
    style.font.name = S.FONT_BODY_LATIN
    style.font.size = Pt(S.SIZE_BODY)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), S.FONT_BODY_CJK)


def _add_title_page(doc, meta: ThesisMeta):
    """标题页：标题居中，作者、单位、日期。"""
    # 顶部空行
    _add_paragraph(doc, "")
    _add_paragraph(doc, "")

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(24)
    run = title_p.add_run(meta.title)
    _set_run_font(run, size=S.SIZE_TITLE, bold=True, cjk=S.FONT_HEADING_CJK)

    if meta.author:
        _add_paragraph(doc, meta.author, align="center")
    if meta.affiliation:
        _add_paragraph(doc, meta.affiliation, align="center")
    if meta.date:
        _add_paragraph(doc, meta.date, align="center")

    if meta.abstract:
        _add_paragraph(doc, "")
        _add_paragraph(doc, "摘要", level=2, align="center")
        _add_paragraph_with_inline_format(doc, meta.abstract, first_line_indent=True)

    if meta.keywords:
        kw_p = doc.add_paragraph()
        kw_p.paragraph_format.first_line_indent = Cm(S.FIRST_LINE_INDENT_CM)
        kw_p.paragraph_format.line_spacing = S.LINE_SPACING
        kw_label = kw_p.add_run("关键词：")
        _set_run_font(kw_label, size=S.SIZE_BODY, bold=True)
        kw_text = kw_p.add_run("、".join(meta.keywords))
        _set_run_font(kw_text, size=S.SIZE_BODY)


def build_thesis_with_custom_cover(
    cover_template_path: str,
    meta: ThesisMeta,
    method_md: str = "",
    result_md: str = "",
    descriptive_table: Optional[pd.DataFrame] = None,
    extra_tables: Optional[List[Tuple[str, pd.DataFrame]]] = None,
    figures: Optional[List[FigureItem]] = None,
    defense_qa_md: str = "",
) -> bytes:
    """v2.8: 用自定义 docx 封面模板 + 系统正文，生成完整论文。

    把模板文件中的所有内容保留作为「封面页」，
    然后在末尾拼接：摘要、关键词、方法、结果、表格、图表、答辩附录。

    Args:
        cover_template_path: 用户上传的 .docx 模板路径
        meta: 摘要/关键词等元信息（标题不再单独渲染——假设模板已含）
        其他参数同 build_thesis_docx

    Returns:
        .docx 字节流；模板格式异常时抛 ValueError。
    """
    try:
        doc = Document(cover_template_path)
    except Exception as e:
        raise ValueError(f"无法读取封面模板（请确认是有效的 .docx 文件）：{e}")

    # 不再调用 _setup_page（保留模板的页边距/页脚/页眉）
    # 不再调用 _add_title_page（模板自带封面）

    # 在模板末尾插入分页符，再追加正文
    doc.add_page_break()

    # 摘要 + 关键词（如果模板里没有摘要，系统补一份）
    if meta.abstract:
        _add_paragraph(doc, "摘要", level=2, align="center")
        _add_paragraph_with_inline_format(doc, meta.abstract, first_line_indent=True)
    if meta.keywords:
        kw_p = doc.add_paragraph()
        kw_p.paragraph_format.first_line_indent = Cm(S.FIRST_LINE_INDENT_CM)
        kw_p.paragraph_format.line_spacing = S.LINE_SPACING
        kw_label = kw_p.add_run("关键词：")
        _set_run_font(kw_label, size=S.SIZE_BODY, bold=True)
        kw_text = kw_p.add_run("、".join(meta.keywords))
        _set_run_font(kw_text, size=S.SIZE_BODY)

    if method_md:
        doc.add_page_break()
        _render_markdown(doc, method_md)
    if result_md:
        _render_markdown(doc, result_md)

    table_no = 1
    if descriptive_table is not None and not descriptive_table.empty:
        _add_dataframe_table(doc, descriptive_table, caption="描述性统计结果", number=table_no)
        table_no += 1
    if extra_tables:
        for caption, df in extra_tables:
            if df is not None and not df.empty:
                _add_dataframe_table(doc, df, caption=caption, number=table_no)
                table_no += 1

    if figures:
        for i, fig in enumerate(figures, start=1):
            _add_figure(doc, fig, number=i)

    if defense_qa_md:
        doc.add_page_break()
        _add_paragraph(doc, "附录 A：答辩问题预演", level=1)
        _render_markdown(doc, defense_qa_md)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# v2.9: 占位符注册表（每个 placeholder_id → 补全指导）
PROTOCOL_PLACEHOLDER_GUIDE: Dict[str, Dict[str, str]] = {
    "research_questions": {
        "section": "一、实验设计概述 → （二）研究问题",
        "guide": "建议补充：研究的核心疑问，应为可被实证检验的问题（≥1 个）",
    },
    "hypotheses": {
        "section": "一、实验设计概述 → （三）研究假设",
        "guide": "建议补充：明确的方向性或非方向性假设（如「H1：实验组 X 高于对照组」）",
    },
    "independent_vars": {
        "section": "一、实验设计概述 → （四）变量",
        "guide": "建议补充：自变量名称、操纵方式、水平数（如「图片情绪：积极/消极」）",
    },
    "dependent_vars": {
        "section": "一、实验设计概述 → （四）变量",
        "guide": "建议补充：因变量名称、测量工具、单位（如「正确率（百分比）」）",
    },
    "inclusion_criteria": {
        "section": "二、被试招募 → （二）纳入标准",
        "guide": "建议补充：年龄、利手、视力、语言能力、健康状况等（≥3 项）",
    },
    "exclusion_criteria": {
        "section": "二、被试招募 → （三）排除标准",
        "guide": "建议补充：精神疾病史、药物使用、近期参与类似研究等（≥2 项）",
    },
    "materials": {
        "section": "三、实验材料 → （一）刺激材料",
        "guide": "应包含：呈现时间、视觉角度、亮度、语音强度、刺激来源（如 IAPS 编号）等",
    },
    "apparatus": {
        "section": "三、实验材料 → （二）设备清单",
        "guide": "建议补充：显示器型号/尺寸、键盘/反应箱、生理记录仪等",
    },
    "procedure": {
        "section": "四、实验流程",
        "guide": "建议补充：每个 trial 的具体顺序，按拉丁方/被试间/被试内展开",
    },
    "ethics": {
        "section": "六、注意事项 → （一）伦理与保密",
        "guide": "建议补充：伦理委员会编号、知情同意书来源、数据匿名化方案",
    },
    "analysis_plan": {
        "section": "六、注意事项 → （三）数据分析方案",
        "guide": "建议补充：统计方法、显著性水平、效应量报告标准、缺失值处理",
    },
}


def _detect_placeholders(design) -> List[Dict[str, str]]:
    """v2.9: 检测设计对象中"待补充"的占位项，返回清单。

    每个返回项含 section / placeholder_id / guide。
    """
    issues: List[Dict[str, str]] = []

    def _is_empty(val) -> bool:
        if val is None:
            return True
        if isinstance(val, str):
            return not val.strip()
        if hasattr(val, "__len__"):
            return len(val) == 0
        return False

    def _check(field: str):
        v = getattr(design, field, None)
        if _is_empty(v):
            info = PROTOCOL_PLACEHOLDER_GUIDE.get(field, {})
            issues.append({
                "field": field,
                "section": info.get("section", field),
                "guide": info.get("guide", "请补充该字段"),
            })

    for field in PROTOCOL_PLACEHOLDER_GUIDE.keys():
        _check(field)

    return issues


def count_protocol_placeholders(design) -> int:
    """v2.9: 给 UI 用 — 仅返回待补项数量（不渲染文档）。"""
    return len(_detect_placeholders(design))


def build_experiment_protocol_docx(design, *, researcher: str = "", date: str = "") -> bytes:
    """v2.8 + v2.9: 生成实验程序文档（.docx）。

    v2.9 新增：检测「待补充」占位项，文档末尾自动追加清单（≥3 个时显示）。

    Args:
        design: ExperimentDesign 数据类实例
        researcher: 研究者姓名
        date: 日期字符串（默认当前日期）

    Returns:
        .docx 字节流
    """
    from datetime import datetime
    if not date:
        date = datetime.now().strftime("%Y 年 %m 月 %d 日")

    # v2.9: 提前检测占位
    placeholders = _detect_placeholders(design)

    doc = Document()
    _setup_page(doc)

    def _v(name: str, default=""):
        return getattr(design, name, default) or default

    # ============ 标题页 ============
    _add_paragraph(doc, "")
    _add_paragraph(doc, "")
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(20)
    run = title_p.add_run(_v("title", "实验程序文档"))
    _set_run_font(run, size=S.SIZE_TITLE, bold=True, cjk=S.FONT_HEADING_CJK)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("实验程序文档（Experiment Protocol）")
    _set_run_font(sub_run, size=S.SIZE_H2, italic=True)

    _add_paragraph(doc, "", level=0)
    if researcher:
        _add_paragraph(doc, f"研究者：{researcher}", align="center")
    _add_paragraph(doc, f"目标群体：{_v('target_population', '未指定')}", align="center")
    _add_paragraph(doc, f"日期：{date}", align="center")

    doc.add_page_break()

    # ============ 一、实验设计概述 ============
    _add_paragraph(doc, "一、实验设计概述", level=1)
    _add_paragraph(doc, "（一）设计类型", level=2)
    _add_paragraph_with_inline_format(doc, _v("design_type_zh", _v("design_type", "（未指定）")))

    _add_paragraph(doc, "（二）研究问题", level=2)
    rqs = _v("research_questions", []) or ["（待补充）"]
    for i, rq in enumerate(rqs, 1):
        _add_paragraph_with_inline_format(doc, f"{i}. {rq}", first_line_indent=False)

    _add_paragraph(doc, "（三）研究假设", level=2)
    hyps = _v("hypotheses", []) or ["（待补充）"]
    for i, h in enumerate(hyps, 1):
        _add_paragraph_with_inline_format(doc, f"H{i}: {h}", first_line_indent=False)

    _add_paragraph(doc, "（四）变量", level=2)
    ivs = _v("independent_vars", [])
    dvs = _v("dependent_vars", [])
    ctrls = _v("control_vars", [])
    iv_text = "、".join(str(x) for x in ivs) if ivs else "（待指定）"
    dv_text = "、".join(str(x) for x in dvs) if dvs else "（待指定）"
    ctrl_text = "、".join(str(x) for x in ctrls) if ctrls else "无"
    _add_paragraph_with_inline_format(doc, f"自变量（IV）：{iv_text}")
    _add_paragraph_with_inline_format(doc, f"因变量（DV）：{dv_text}")
    _add_paragraph_with_inline_format(doc, f"控制变量：{ctrl_text}")

    # ============ 二、被试招募 ============
    doc.add_page_break()
    _add_paragraph(doc, "二、被试招募", level=1)

    _add_paragraph(doc, "（一）样本量依据", level=2)
    n_subjects = _v("n_subjects", 0)
    n_per_group = _v("n_per_group", 0)
    n_groups = _v("n_groups", 0)
    sample_text = (
        f"本实验计划招募 {n_subjects} 名被试"
        + (f"（{n_groups} 组，每组 {n_per_group} 人）。" if n_groups else "。")
        + " 样本量基于 G*Power 软件计算（α=.05, power=.80, 预期中等效应量）确定。"
    )
    _add_paragraph_with_inline_format(doc, sample_text)

    _add_paragraph(doc, "（二）纳入标准", level=2)
    incs = _v("inclusion_criteria", []) or ["（待补充）"]
    for c in incs:
        _add_paragraph_with_inline_format(doc, f"• {c}", first_line_indent=False)

    _add_paragraph(doc, "（三）排除标准", level=2)
    excs = _v("exclusion_criteria", []) or ["（待补充）"]
    for c in excs:
        _add_paragraph_with_inline_format(doc, f"• {c}", first_line_indent=False)

    _add_paragraph(doc, "（四）知情同意", level=2)
    _add_paragraph_with_inline_format(doc, (
        "所有被试在参与前需签署知情同意书。同意书包含：研究目的、实验流程、"
        "可能的风险与收益、数据保密措施、自愿原则（可随时退出）、研究者联系方式。"
    ))

    # ============ 三、实验材料 ============
    doc.add_page_break()
    _add_paragraph(doc, "三、实验材料", level=1)

    _add_paragraph(doc, "（一）刺激材料", level=2)
    materials = _v("materials", []) or []
    if materials:
        for i, m in enumerate(materials, 1):
            if isinstance(m, dict):
                name = m.get("name", f"材料{i}")
                desc = m.get("description", "")
                _add_paragraph_with_inline_format(doc, f"{i}. **{name}**：{desc}")
            else:
                _add_paragraph_with_inline_format(doc, f"{i}. {m}")
    else:
        _add_paragraph_with_inline_format(doc, "（请填写实验中使用的刺激材料、问卷、量表名称及来源）")

    _add_paragraph(doc, "（二）设备清单", level=2)
    apparatus = _v("apparatus", []) or []
    if apparatus:
        for a in apparatus:
            _add_paragraph_with_inline_format(doc, f"• {a}", first_line_indent=False)
    else:
        _add_paragraph_with_inline_format(doc, "（待补充：电脑、显示器规格、耳机、生理记录仪等）")

    # ============ 四、实验流程 ============
    doc.add_page_break()
    _add_paragraph(doc, "四、实验流程", level=1)

    procedure = _v("procedure", "")
    if procedure:
        if isinstance(procedure, list):
            for i, step in enumerate(procedure, 1):
                _add_paragraph_with_inline_format(doc, f"{i}. {step}")
        else:
            _add_paragraph_with_inline_format(doc, str(procedure))
    else:
        steps = [
            "被试到达实验室，研究者介绍实验目的并签署知情同意书。",
            "被试填写人口学信息问卷。",
            "正式实验任务（具体每个 trial 顺序按拉丁方/被试间/被试内分配）。",
            "实验后访谈（manipulation check）。",
            "向被试致谢并发放报酬。",
        ]
        for i, s in enumerate(steps, 1):
            _add_paragraph_with_inline_format(doc, f"{i}. {s}")

    # ============ 五、数据记录字段 ============
    doc.add_page_break()
    _add_paragraph(doc, "五、数据记录字段", level=1)
    _add_paragraph_with_inline_format(doc, "实验数据应至少记录以下字段：")
    fields = [
        "被试编号（subject_id）",
        "组别 / 实验条件",
        f"自变量取值：{iv_text}",
        f"因变量数值：{dv_text}",
        "反应时（如有）",
        "正确率（如有）",
        "刺激呈现时间戳",
        "其他备注",
    ]
    for f in fields:
        _add_paragraph_with_inline_format(doc, f"• {f}", first_line_indent=False)

    # ============ 六、注意事项 ============
    doc.add_page_break()
    _add_paragraph(doc, "六、注意事项", level=1)

    _add_paragraph(doc, "（一）伦理与保密", level=2)
    ethics = _v("ethics", "")
    if ethics:
        _add_paragraph_with_inline_format(doc, ethics)
    else:
        _add_paragraph_with_inline_format(doc, (
            "本研究遵循《赫尔辛基宣言》，所有数据匿名化处理，"
            "原始数据仅由研究团队接触，发表论文时不出现可识别个人信息。"
        ))

    _add_paragraph(doc, "（二）突发情况处理", level=2)
    _add_paragraph_with_inline_format(doc, (
        "若被试中途身体或情绪不适，立即终止实验并允许退出，"
        "已收集的数据按被试要求处理（保留或删除）。"
        "实验过程中如设备故障，记录故障时间并联系技术支持。"
    ))

    _add_paragraph(doc, "（三）数据分析方案", level=2)
    plan = _v("analysis_plan", "")
    plan_detailed = _v("analysis_plan_detailed", "")
    if plan_detailed:
        _add_paragraph_with_inline_format(doc, plan_detailed)
    elif plan:
        _add_paragraph_with_inline_format(doc, plan)
    else:
        _add_paragraph_with_inline_format(doc, "（待补充统计方法、显著性水平、效应量报告标准）")

    # ============ 末尾备注 ============
    notes = _v("notes", "")
    if notes:
        _add_paragraph(doc, "（四）补充说明", level=2)
        _add_paragraph_with_inline_format(doc, notes)

    # ============ v2.9: 待补充事项清单（≥3 个时显示）============
    if len(placeholders) >= 3:
        doc.add_page_break()
        _add_paragraph(doc, "📝 待补充事项清单", level=1)
        _add_paragraph_with_inline_format(doc, (
            f"系统检测到本文档共有 {len(placeholders)} 处「待补充」内容，"
            f"建议在正式实验前逐项填写完整。完整的实验程序文档应作为预注册的附件，"
            f"也是答辩时评委可能要求展示的材料。"
        ))
        _add_paragraph(doc, "")

        for i, item in enumerate(placeholders, start=1):
            line_p = doc.add_paragraph()
            line_p.paragraph_format.line_spacing = S.LINE_SPACING
            line_p.paragraph_format.space_after = Pt(4)
            num_run = line_p.add_run(f"{i}. ")
            _set_run_font(num_run, size=S.SIZE_BODY, bold=True)
            sec_run = line_p.add_run(f"{item['section']}")
            _set_run_font(sec_run, size=S.SIZE_BODY, bold=True)

            guide_p = doc.add_paragraph()
            guide_p.paragraph_format.left_indent = Cm(0.74)
            guide_p.paragraph_format.line_spacing = S.LINE_SPACING
            guide_p.paragraph_format.space_after = Pt(8)
            guide_run = guide_p.add_run(f"   ▸ {item['guide']}")
            _set_run_font(guide_run, size=S.SIZE_CAPTION, italic=True)

        # 完成度提示
        _add_paragraph(doc, "")
        tip_p = doc.add_paragraph()
        tip_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tip_run = tip_p.add_run(
            "💡 提示：所有「待补充」项填写完毕后，可重新生成本文档，清单将自动消失。"
        )
        _set_run_font(tip_run, size=S.SIZE_CAPTION, italic=True, color_hex="888888")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_thesis_docx(
    meta: ThesisMeta,
    method_md: str = "",
    result_md: str = "",
    descriptive_table: Optional[pd.DataFrame] = None,
    extra_tables: Optional[List[Tuple[str, pd.DataFrame]]] = None,
    figures: Optional[List[FigureItem]] = None,
    defense_qa_md: str = "",
) -> bytes:
    """构建完整的 .docx 论文初稿。

    Args:
        meta: 标题/作者/摘要等元信息
        method_md: 方法部分 Markdown
        result_md: 结果部分 Markdown
        descriptive_table: 描述统计表（自动作为表1）
        extra_tables: [(caption, df), ...] 追加表格（表2,3,...）
        figures: [FigureItem,...] 论文版图表（嵌入到结果之后）
        defense_qa_md: 答辩问题部分 Markdown（可选，附录）

    Returns:
        .docx 字节流
    """
    doc = Document()
    _setup_page(doc)
    _add_title_page(doc, meta)

    if method_md:
        doc.add_page_break()
        _render_markdown(doc, method_md)

    if result_md:
        _render_markdown(doc, result_md)

    table_no = 1
    if descriptive_table is not None and not descriptive_table.empty:
        _add_dataframe_table(doc, descriptive_table, caption="描述性统计结果", number=table_no)
        table_no += 1

    if extra_tables:
        for caption, df in extra_tables:
            if df is not None and not df.empty:
                _add_dataframe_table(doc, df, caption=caption, number=table_no)
                table_no += 1

    if figures:
        for i, fig in enumerate(figures, start=1):
            _add_figure(doc, fig, number=i)

    if defense_qa_md:
        doc.add_page_break()
        _add_paragraph(doc, "附录 A：答辩问题预演", level=1)
        _render_markdown(doc, defense_qa_md)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# 问卷正式文档（v4.1）
# --------------------------------------------------------------------------- #

DEFAULT_LIKERT_ANCHORS = {
    5: ["完全不符合", "比较不符合", "一般", "比较符合", "完全符合"],
    7: ["完全不符合", "不符合", "比较不符合", "一般", "比较符合", "符合", "完全符合"],
}


def _add_likert_table(doc, items: list, reverse_set: set, scale_points: int):
    """Likert 答题表：列 = 题号 / 题目 / 1..N。"""
    n_rows = len(items) + 1
    n_cols = 2 + scale_points

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    headers = ["题号", "题目"] + [str(i + 1) for i in range(scale_points)]
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        _set_run_font(run, size=S.SIZE_TABLE, bold=True, cjk=S.FONT_HEADING_CJK)

    for i, item in enumerate(items, start=1):
        is_rev = (i - 1) in reverse_set
        c0 = table.rows[i].cells[0]
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        num_text = f"{i}" + (" (R)" if is_rev else "")
        run = c0.paragraphs[0].add_run(num_text)
        _set_run_font(run, size=S.SIZE_TABLE, bold=is_rev)

        c1 = table.rows[i].cells[1]
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = c1.paragraphs[0].add_run(str(item))
        _set_run_font(run, size=S.SIZE_TABLE)

        for k in range(scale_points):
            c = table.rows[i].cells[2 + k]
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = c.paragraphs[0].add_run("○")
            _set_run_font(run, size=S.SIZE_TABLE)

    _apply_three_line_borders(table)


def build_questionnaire_docx(
    items_doc,
    *,
    scale_points: int = 5,
    anchors: Optional[List[str]] = None,
    header_meta: Optional[dict] = None,
    show_id_field: bool = True,
) -> bytes:
    """构建正式问卷 Word 文档（v4.1）。

    Args:
        items_doc: src.questionnaire.items_loader.ItemsDoc 或具相同字段的对象
        scale_points: Likert 点数（建议 5 或 7）
        anchors: 自定义锚点（长度需等于 scale_points）；不传走默认中文锚点
        header_meta: 可选 dict，含 researcher / project / version / date
        show_id_field: 是否打印「编号 / 性别 / 年龄 / 日期」填写区

    Returns:
        .docx 字节流
    """
    if scale_points < 2 or scale_points > 11:
        raise ValueError(f"scale_points 需在 2-11 之间（传入 {scale_points}）")

    if anchors is None:
        anchors = DEFAULT_LIKERT_ANCHORS.get(
            scale_points,
            [str(i + 1) for i in range(scale_points)],
        )
    if len(anchors) != scale_points:
        raise ValueError(
            f"anchors 长度 ({len(anchors)}) 需与 scale_points ({scale_points}) 一致"
        )

    title = (getattr(items_doc, "title", "") or "心理测量问卷").strip()
    instructions = (getattr(items_doc, "instructions", "") or "").strip()
    items = list(getattr(items_doc, "items", []) or [])
    reverse_set = set(getattr(items_doc, "reverse_indices", []) or [])

    if not items:
        raise ValueError("ItemsDoc.items 为空，无法生成问卷文档。")

    doc = Document()
    _setup_page(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(title)
    _set_run_font(
        run, size=S.SIZE_TITLE, bold=True,
        latin=S.FONT_HEADING_LATIN, cjk=S.FONT_HEADING_CJK,
    )

    if header_meta:
        meta_lines = []
        for key, label in [
            ("researcher", "主试"),
            ("project", "研究项目"),
            ("version", "版本"),
            ("date", "日期"),
        ]:
            val = header_meta.get(key)
            if val:
                meta_lines.append(f"{label}：{val}")
        if meta_lines:
            mp = doc.add_paragraph()
            mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = mp.add_run("    ".join(meta_lines))
            _set_run_font(run, size=S.SIZE_CAPTION)

    if show_id_field:
        idp = doc.add_paragraph()
        idp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        idp.paragraph_format.space_before = Pt(6)
        run = idp.add_run("编号：__________    性别：____    年龄：____    日期：__________")
        _set_run_font(run, size=S.SIZE_BODY)

    if instructions:
        _add_paragraph(doc, "指导语", level=2)
        _add_paragraph_with_inline_format(doc, instructions, first_line_indent=True)

    _add_paragraph(doc, "评分说明", level=2)
    anchor_lines = [f"{i + 1} = {anchors[i]}" for i in range(scale_points)]
    legend_text = (
        "请在每道题后选择最符合您实际情况的数字（圈选）："
        + "；".join(anchor_lines) + "。"
    )
    _add_paragraph_with_inline_format(doc, legend_text, first_line_indent=False)
    if reverse_set:
        rev_note = f"题号后标 (R) 的为反向题，共 {len(reverse_set)} 题，计分时需反向。"
        _add_paragraph_with_inline_format(doc, rev_note, first_line_indent=False)

    _add_paragraph(doc, "题项", level=2)
    _add_likert_table(doc, items, reverse_set, scale_points)

    doc.add_paragraph()
    thank_p = doc.add_paragraph()
    thank_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    thank_p.paragraph_format.space_before = Pt(12)
    run = thank_p.add_run("—— 问卷结束，感谢您的参与 ——")
    _set_run_font(run, size=S.SIZE_BODY, bold=True)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
